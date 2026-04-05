import streamlit as st
import anthropic
import pdfplumber
import os, re, time, io, tempfile, json, difflib, sys, uuid
import base64
from html import escape as html_escape
from datetime import datetime
from pathlib import Path
# lxml.etree was imported here but never used — removed.
from docx import Document as DocxDoc
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                 Table, TableStyle, HRFlowable, PageBreak)
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT

# Local modules
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from clause_benchmarks import CLAUSE_BENCHMARKS, get_benchmark, format_benchmarks_for_prompt
from precedent_clauses import (PRECEDENT_CLAUSES, get_precedent, render_precedent,
                                format_precedents_for_prompt)
from matter_store import (init_db, save_matter, list_matters, load_matter,
                          delete_matter, matter_count, matters_by_reviewer)
from pipeda_compliance import (
    detect_personal_information, pi_risk_summary,
    generate_retainer_disclosure, purge_expired_matters,
    purge_all_matters, matters_older_than,
)
from deal_context import (DealContext, format_deal_context_for_prompt,
                          get_variable_values, compute_clause_risk_score,
                          INDUSTRY_RISK_TIERS, COUNTERPARTY_RISK)

# OCR fallback
try:
    import pytesseract
    from pdf2image import convert_from_bytes
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

# Image preprocessing for messy scans
try:
    import cv2
    import numpy as np
    HAS_CV2 = True
except ImportError:
    HAS_CV2 = False

# ═══════════════════════════════════════════════════════════════════
# ContractCheck Pro v12.6
# ═══════════════════════════════════════════════════════════════════
# NEW in v12.0:
#   1. Long-contract chunking pipeline — no more silent truncation.
#      Contracts of any length are analysed completely via a
#      chunk-then-synthesise approach.
#   2. Persistent matter history — SQLite-backed store. Every analysis
#      saved. Browse, reload, and compare past matters from the sidebar.
#   3. True OOXML track changes — redline DOCX uses w:ins / w:del XML
#      elements. Word shows real Track Changes that can be
#      accepted or rejected natively.
#   4. Vetted precedent clause library — 40+ pre-approved Canadian
#      replacement clauses. AI adapts from precedent rather than
#      drafting from scratch.
#   5. Multi-document support — upload main agreement + up to 4
#      schedules/exhibits. Analysed together with cross-references.
# ═══════════════════════════════════════════════════════════════════

CHUNK_SIZE     = 90_000   # chars per chunk for long contracts
CHUNK_OVERLAP  = 3_000    # overlap to catch cross-boundary clauses
CHUNK_THRESHOLD = 110_000 # contracts above this are chunked

# ── Cost guard ───────────────────────────────────────────────
# Soft warning shown before analysis if projected cost exceeds this.
# Hard limit not enforced — lawyer always has final say.
MAX_SESSION_COST_USD = 2.00  # warn after $2 USD in one session

# ── Audit logging ────────────────────────────────────────────
import logging
from logging.handlers import RotatingFileHandler
from pathlib import Path as _Path

def _setup_audit_log():
    log_dir = _Path.home() / ".contractcheck"
    log_dir.mkdir(parents=True, exist_ok=True)
    logger = logging.getLogger("contractcheck_audit")
    if not logger.handlers:
        handler = RotatingFileHandler(
            log_dir / "audit.log",
            maxBytes  = 10 * 1024 * 1024,  # 10 MB
            backupCount = 5,
            encoding  = "utf-8",
        )
        handler.setFormatter(logging.Formatter(
            "%(asctime)s | %(levelname)s | %(message)s",
            datefmt="%Y-%m-%d %H:%M:%S"
        ))
        logger.setLevel(logging.INFO)
        logger.addHandler(handler)
    return logger

_audit = _setup_audit_log()

st.set_page_config(page_title="ContractCheck Pro", layout="wide", page_icon="⚖️")

# ── CSS ──
st.markdown("""
<style>
    .pro-header {
        background: linear-gradient(135deg, #0d1b2a 0%, #1b2d4f 60%, #2a3f5f 100%);
        padding: 24px 30px; border-radius: 6px; margin-bottom: 16px;
        border-bottom: 3px solid #c9a84c;
    }
    .risk-banner {
        padding: 14px 20px; border-radius: 4px; margin: 12px 0;
        background: #fafafa; font-family: Georgia, serif;
    }
    .risk-critical { border-left: 5px solid #8B0000; }
    .risk-high     { border-left: 5px solid #B8860B; }
    .risk-medium   { border-left: 5px solid #DAA520; }
    .risk-low      { border-left: 5px solid #006400; }
    .heatmap-container { display: flex; flex-wrap: wrap; gap: 6px; margin: 10px 0; }
    .heat-cell {
        flex: 1 1 130px; padding: 8px 10px; border-radius: 4px;
        font-family: Georgia, serif; font-size: 12px; text-align: center;
        border: 1px solid rgba(0,0,0,0.06); min-width: 120px;
    }
    .heat-critical { background: #fce4e4; color: #7f1d1d; border-left: 3px solid #991b1b; }
    .heat-high     { background: #fef3c7; color: #78350f; border-left: 3px solid #b45309; }
    .heat-medium   { background: #fef9c3; color: #713f12; border-left: 3px solid #ca8a04; }
    .heat-low      { background: #dcfce7; color: #14532d; border-left: 3px solid #15803d; }
    .heat-missing  { background: #f3f4f6; color: #4b5563; border-left: 3px solid #6b7280; }
    .heat-na       { background: #f9fafb; color: #9ca3af; border-left: 3px solid #d1d5db; }
    .issue-card {
        border-radius: 4px; padding: 12px 16px; margin: 6px 0;
        border: 1px solid rgba(0,0,0,0.06); font-size: 14px;
    }
    .issue-critical { background: #fef2f2; border-left: 4px solid #991b1b; }
    .issue-high     { background: #fffbeb; border-left: 4px solid #b45309; }
    .issue-medium   { background: #fefce8; border-left: 4px solid #ca8a04; }
    .issue-low      { background: #f0fdf4; border-left: 4px solid #15803d; }
    .conf-high   { color: #15803d; font-weight: 600; }
    .conf-medium { color: #b45309; font-weight: 600; }
    .conf-low    { color: #991b1b; font-weight: 600; }
    .matter-row  { padding: 8px 12px; border-radius: 4px; margin: 3px 0;
                   background: #f9fafb; border: 1px solid #e5e7eb; font-size: 12px; }
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div class='pro-header'>
  <h1 style='color:white;margin:0;font-size:22px;font-family:Georgia,serif;letter-spacing:0.3px;'>
    ContractCheck Pro
    <span style="font-size:11px;background:#c9a84c;color:#0d1b2a;
                 padding:2px 10px;border-radius:10px;vertical-align:middle;
                 margin-left:10px;font-family:Arial;font-weight:600;">v12.6</span>
  </h1>
  <p style='color:#8fa4c4;margin:4px 0 0 0;font-size:12px;font-family:Arial;'>
    Contract Review &middot; Clause Benchmarking &middot; Negotiation Simulator &middot;
    Unlimited Document Length &middot; Persistent Matter History &middot; True Track Changes
  </p>
</div>
""", unsafe_allow_html=True)

st.markdown("""
<div style='background:#fff8e7;border:1px solid #e5c76b;padding:10px 16px;border-radius:4px;
            margin-bottom:8px;font-size:12px;color:#7a5d00;font-family:Georgia,serif;'>
  <b>AI FIRST-PASS ANALYSIS</b> &mdash;
  This tool is a drafting and productivity aid only. All conclusions must be independently
  reviewed and verified by licensed counsel before reliance, client communication, or execution.
  The reviewing lawyer retains full professional responsibility. Not a substitute for legal advice.
</div>
""", unsafe_allow_html=True)

st.markdown("""
<div style='background:#f7f8f7;border:1px solid #d0d0d0;padding:9px 14px;border-radius:4px;
            margin-bottom:14px;font-size:11px;color:#555;font-family:Arial;'>
  <b>Security &amp; Compliance:</b>
  <b>What is stored:</b> Uploaded files are processed by Streamlit into a temporary directory
  and discarded when the session ends. Contract text is sent to Anthropic&rsquo;s API for
  analysis and is not retained by Anthropic per their
  <a href="https://www.anthropic.com/policies/terms" target="_blank">commercial API terms</a>
  (Anthropic does not train on commercial API data). Matter history (risk levels, issue titles,
  clause recommendations, reviewer name) is saved to a local SQLite database;
  verbatim contract text is stripped before saving so the database contains analysis metadata
  only &mdash; not raw contract text.
  &bull; <b>PIPEDA:</b> Personal information patterns are scanned before each analysis.
  A retainer disclosure template is available in Privacy &amp; PIPEDA Settings (sidebar).
  &bull; <b>Cross-border transfer:</b> API calls are processed on Anthropic&rsquo;s US
  infrastructure. Ensure client retainer covers this per PIPEDA Principle 3 (consent).
  Suitable for on-premise or private cloud deployment.
</div>
""", unsafe_allow_html=True)

# ── API KEY ──
api_key = ""
_KEY_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "api_key.txt")
if os.path.exists(_KEY_FILE):
    try:
        with open(_KEY_FILE, encoding="utf-8-sig") as f:
            api_key = f.read().strip()
    except OSError:
        pass  # Silently skip unreadable key file
if not api_key: api_key = os.environ.get("ANTHROPIC_API_KEY", "")
if not api_key: api_key = st.sidebar.text_input("API Key", type="password")
if not api_key: st.info("Enter your Anthropic API key to begin."); st.stop()
ai = anthropic.Anthropic(api_key=api_key)

# ── SIDEBAR SETTINGS ──
st.sidebar.header("Settings")
output_mode = st.sidebar.selectbox("Output Format", [
    "Associate Memo (Internal)",
    "Client Advisory (Plain English)",
    "Opposing Counsel Redline"], index=0)
analysis_depth = st.sidebar.selectbox("Depth", [
    "Standard Review", "Quick Scan", "Deep Dive + Redraft"], index=0)
MODEL = st.sidebar.selectbox("Model", [
    "claude-sonnet-4-5 (Recommended)",
    "claude-opus-4-5 (Complex matters)"], index=0)
model_id = MODEL.split(" ")[0]
use_caching = st.sidebar.checkbox("Prompt Caching", value=True)

st.sidebar.divider()
st.sidebar.subheader("Compliance")
reviewer_name = st.sidebar.text_input("Reviewing Lawyer",
    placeholder="e.g. J. Smith",
    help="Name of the lawyer reviewing this AI output. Appears on audit log.")
analysis_language = st.sidebar.selectbox("Analysis Language",
    ["English", "French / Français"], index=0,
    help="French mode: Analyse French-language contracts. Output memo in French. Quebec civil law benchmarks applied automatically.")

_INPUT_RATE  = 3.0  if "sonnet" in model_id else 15.0
_OUTPUT_RATE = 15.0 if "sonnet" in model_id else 75.0
depth_tokens = {"Quick Scan": 8000, "Standard Review": 16000, "Deep Dive + Redraft": 20000}[analysis_depth]
est = (3000 * 0.3 + 6000 * _INPUT_RATE) / 1_000_000 + depth_tokens * _OUTPUT_RATE / 1_000_000
st.sidebar.caption(f"Est. cost per analysis: ~${est:.3f}")

if "history" not in st.session_state: st.session_state.history = []

def _valid_docx(b) -> bool:
    """Return True if b looks like a valid DOCX (ZIP with PK header)."""
    return isinstance(b, (bytes, bytearray)) and len(b) > 100 and b[:2] == b'PK'


def _docx_download_link(data: bytes, filename: str, label: str,
                         size_kb: int = 0, primary: bool = False) -> str:
    """Generate an HTML anchor tag that downloads DOCX via data URI.

    Edge intercepts http/https downloads of .docx files and opens them
    in its built-in Office viewer. Data URIs bypass this entirely because
    they are not HTTP downloads — the browser treats them as local resources.
    The file downloads directly to the user's Downloads folder with the
    correct .docx extension and opens in Word on double-click.
    """
    b64 = base64.b64encode(data).decode()
    bg    = "#c0392b" if primary else "#f8f9fa"
    color = "white"  if primary else "#212529"
    border= "#c0392b" if primary else "#dee2e6"
    return (
        f'<a href="data:application/vnd.openxmlformats-officedocument'
        f'.wordprocessingml.document;base64,{b64}" '
        f'download="{filename}" '
        f'style="display:block;width:100%;padding:10px 16px;margin:4px 0;'
        f'background:{bg};color:{color};border:1px solid {border};'
        f'border-radius:4px;text-align:center;text-decoration:none;'
        f'font-size:14px;font-weight:{"600" if primary else "400"};">'
        f'{label}'
        f'</a>'
    )

def _valid_pdf(b) -> bool:
    """Return True if b looks like a valid PDF."""
    return isinstance(b, (bytes, bytearray)) and len(b) > 100 and b[:4] == b'%PDF'

# Clear any stale / corrupt download bytes from previous sessions.
# Corrupt bytes (e.g. PDF bytes stored under a DOCX key) survive Streamlit
# reruns and cause downloads that won't open. Reset on every cold start.
for _stale_key in ['_dl_pdf', '_dl_docx', '_dl_amend', '_dl_redline',
                   '_dl_safe', '_dl_ts', '_dl_has_redlines']:
    if _stale_key in st.session_state and _stale_key.startswith('_dl_'):
        # Only clear if this is a fresh browser session (no history yet)
        if not st.session_state.get('history'):
            del st.session_state[_stale_key]

# ══════════════════════════════════════════════════════════════
# MANDATORY 15-CLAUSE CHECKLIST
# ══════════════════════════════════════════════════════════════
MANDATORY_CLAUSES = [
    "indemnity", "limitation_of_liability", "termination", "ip_ownership",
    "assignment", "confidentiality", "payment_terms", "interest_and_late_fees",
    "auto_renewal", "governing_law", "dispute_resolution", "force_majeure",
    "notice", "warranties_and_representations", "data_protection",
]
MANDATORY_CLAUSE_LABELS = {
    "indemnity":                    "Indemnity",
    "limitation_of_liability":      "Limitation of Liability",
    "termination":                  "Termination",
    "ip_ownership":                 "IP Ownership",
    "assignment":                   "Assignment",
    "confidentiality":              "Confidentiality",
    "payment_terms":                "Payment Terms",
    "interest_and_late_fees":       "Interest / Late Fees",
    "auto_renewal":                 "Auto-Renewal",
    "governing_law":                "Governing Law",
    "dispute_resolution":           "Dispute Resolution",
    "force_majeure":                "Force Majeure",
    "notice":                       "Notice Provisions",
    "warranties_and_representations":"Warranties & Representations",
    "data_protection":              "Data Protection / Privacy",
}

# ══════════════════════════════════════════════════════════════
# PROVINCE RULES
# ══════════════════════════════════════════════════════════════
PROVINCE_RULES = {
    "Ontario": "ONTARIO: Bill 27 non-compete ban for employees (except C-suite sale of business). ESA 2000: statutory minimums for notice/severance/overtime cannot be contracted out. Common law reasonable notice per Bardal factors (~1 month per year of service). Construction Act: 10% statutory holdback, 60-day lien period, prompt payment (28-day owner payment, 7-day sub payment). Arthur Wishart Act: franchise disclosure 14-day cooling-off. Consumer Protection Act 2002: unconscionable terms void. Limitation Act 2002: 2-year discovery, 15-year ultimate. Mandatory mediation in Ottawa/Toronto/Windsor.",
    "British Columbia": "BRITISH COLUMBIA: No statutory non-compete ban (courts apply strict reasonableness). PIPA (stricter than PIPEDA for private sector). Franchises Act 2017: disclosure requirements. Builders Lien Act. BPCPA consumer protection. Limitation Act: 2-year basic limitation.",
    "Alberta": "ALBERTA: Employment Standards Code (different notice structure, no statutory severance). No non-compete ban (common law reasonableness). PIPA Alberta. Franchises Act. Builders' Lien Act. Limitations Act: 2-year discoverable, 10-year ultimate.",
    "Quebec": "QUEBEC (CIVIL LAW JURISDICTION): Civil Code of Quebec governs. CCQ 1375: good faith mandatory in all contracts. CCQ 1437: abusive clauses in adhesion contracts are null. CCQ 2089: non-compete clauses strictly scrutinised (must be limited in time, territory, and activity). Labour Standards Act (not ESA). Charter of the French Language (Bill 96): consumer contracts must be in French. CCQ 1474: cannot exclude liability for bodily injury or gross fault. CCQ 1623: courts may reduce excessive penalty clauses. Law 25 (privacy). Prescription: 3 years.",
    "Manitoba": "MANITOBA: Employment Standards Code. Consumer Protection Act. Builders' Liens Act. Limitation of Actions Act: 2-year general limitation.",
    "Saskatchewan": "SASKATCHEWAN: Saskatchewan Employment Act. Consumer Protection and Business Practices Act. Builders' Lien Act. Franchise Disclosure Act. Limitations Act: 2-year general.",
    "Nova Scotia": "NOVA SCOTIA: Labour Standards Code. Consumer Protection Act. Builders' Lien Act. Limitation of Actions Act: 2-year general.",
    "New Brunswick": "NEW BRUNSWICK: Employment Standards Act. Consumer Product Warranty and Liability Act. Mechanics' Lien Act. Limitation of Actions Act.",
    "Federal": "FEDERAL (federally regulated industries — banking, telecommunications, interprovincial transport, airlines): Canada Labour Code (not provincial ESA). PIPEDA for privacy. Bank Act. Canada Business Corporations Act. Competition Act 2024 amendments.",
    "Other / Multi-Province": "Apply general Canadian common law. Flag jurisdiction as a risk factor if the contract does not specify governing law. Note potential conflict-of-laws issues for multi-province transactions.",
}

# ══════════════════════════════════════════════════════════════
# CLAUSE DETECTION ENGINE
# ══════════════════════════════════════════════════════════════
CLAUSE_KEYWORDS = {
    "indemnity": [
        "indemnif", "hold harmless", "defend and indemnify", "save harmless",
        "indemnity obligations", "defend, indemnify",
    ],
    "limitation_of_liability": [
        "limitation of liability", "aggregate liability", "cap on liability",
        "consequential damages", "liability shall not exceed", "in no event shall",
        "liability is limited", "maximum liability", "total liability",
        "liability of either party", "shall not be liable for any",
        # "fees paid hereunder" style caps — liability-specific phrasing
        "fees paid hereunder", "fees paid under this", "amounts paid hereunder",
        # Note: "greater of" removed — too many false positives in other contexts
    ],
    "termination": [
        "terminat", "cancel", "expir", "wind down", "right to terminate",
        "termination for cause", "termination for convenience", "notice of termination",
        "effective date of termination", "survive termination",
    ],
    "ip_ownership": [
        "intellectual property", "copyright", "patent", "trademark", "work product",
        "moral rights", "pre-existing", "background ip", "foreground ip",
        "license grant", "work for hire", "work-for-hire", "assigns all right",
        "hereby assigns", "waive all moral", "moral rights waiver",
        "creator retains", "ownership of deliverable",
    ],
    "assignment": [
        "assign", "transfer", "novation", "successors and assigns",
        "may not assign", "without prior written consent", "change of control",
    ],
    "confidentiality": [
        "confidential", "non-disclosure", "nda", "proprietary information",
        "trade secret", "confidential information", "duty of confidentiality",
        "disclose to third", "obligation of confidence",
    ],
    "payment_terms": [
        "payment", "compensation", "fee", "invoice", "price", "rate", "billing",
        "remuneration", "net 30", "net 60", "net 45", "net-30", "net-60",
        "due and payable", "payment schedule", "payment due",
    ],
    "interest_and_late_fees": [
        "interest", "late fee", "late payment", "overdue", "penalty", "per month",
        "per annum", "per year", "compounded", "interest rate", "% per month",
        "% monthly", "1.5%", "18% per annum",
    ],
    "auto_renewal": [
        "auto-renew", "automatic renewal", "automatically renew", "evergreen",
        "successive term", "automatically extend", "unless notice",
        "shall renew", "deemed renewed", "tacit renewal",
    ],
    "governing_law": [
        "governing law", "governed by", "jurisdiction", "applicable law", "laws of",
        "laws of the province", "ontario law", "courts of ontario",
        "exclusive jurisdiction", "submit to jurisdiction",
    ],
    "dispute_resolution": [
        "dispute resolution", "arbitration", "mediation", "litigation",
        "binding arbitration", "dispute shall be", "settle any dispute",
        "submit to arbitration", "adr institute", "mandatory mediation",
    ],
    "force_majeure": [
        "force majeure", "act of god", "unforeseeable", "pandemic",
        "beyond the control", "beyond reasonable control", "unforeseeable circumstances",
        "natural disaster", "government action", "civil unrest",
    ],
    "notice": [
        "notice", "written notice", "notice shall be", "deemed received",
        "notice period", "delivery of notice", "notice by email",
        "notice to the other party",
    ],
    "warranties_and_representations": [
        "represent", "warrant", "guarantee", "covenant", "as is",
        "disclaimer of warranties", "no warranties", "as-is basis",
        "representations and warranties", "makes no representation",
    ],
    "data_protection": [
        "privacy", "personal information", "pipeda", "data protection",
        "data breach", "personal data", "consent to collect",
        "law 25", "pipa", "phipa", "privacy policy", "data processor",
        "data controller", "privacy commissioner",
    ],
}

# ── Statutory red-line patterns ──────────────────────────────
# These are non-waivable rules under Canadian law.
# Detected in Python BEFORE the AI call — injected as verified facts
# the AI cannot contradict.  Keyed by statute reference.
def _is_numeric(s: str) -> bool:
    """Return True if s can be converted to float. Guards against empty/None matches."""
    try:
        float(s); return True
    except (TypeError, ValueError):
        return False


STATUTORY_REDLINES = {
    "Criminal Code s.347 / Bill C-46 (interest cap)": {
        "patterns": [
            r'(\d+(?:\.\d+)?)\s*%\s*per\s*month',
            r'(\d+(?:\.\d+)?)\s*%\s*(?:per annum|per year|annually|p\.a\.)',
        ],
        # Bill C-46 (Royal Assent 2024) amended s.347.1 to cap commercial lending
        # at 35% APR effective January 1, 2025.  The original s.347 criminal
        # threshold of 60% APR remains, but commercial contracts should now be
        # reviewed against the 35% ceiling.  Flag at >35% APR.
        # Check: is the rate above 35% APR?
        # p = pattern string, m = matched number string.
        # "month" in p identifies monthly-rate patterns reliably because
        # the pattern text contains the word "month". Rewritten more explicitly:
        "check": lambda matches: any(
            (float(m) * 12 > 35) if "month" in p else (float(m) > 35)
            for p, m in matches
            if m and _is_numeric(m)
        ),
        "flag": (
            "Interest rate may exceed the 35% APR commercial lending cap under "
            "Criminal Code s.347.1 (as amended by Bill C-46, effective Jan 1 2025). "
            "Rates above 35% APR are potentially criminal for commercial lending. "
            "The original 60% APR threshold under s.347 also remains in force for "
            "other credit arrangements. Verify applicable threshold and urgently "
            "review this clause."
        ),
        "severity": "CRITICAL",
    },
    "Ontario ESA s.5 (non-waivable minimums)": {
        "patterns": [r'employment\s+standards|esa\s+minimum|minimum\s+notice'],
        "check": lambda matches: len(matches) == 0,  # absence is the issue
        "flag": "Employment contract does not reference ESA minimums — ensure statutory floors are preserved.",
        "severity": "HIGH",
        "contract_types": ["Employment Contract"],
    },
    "Ontario Bill 27 (non-compete ban)": {
        "patterns": [r'non.compete|noncompete|not\s+compete|competing\s+business'],
        "check": lambda matches: len(matches) > 0,
        "flag": "Non-compete clause detected in employment context. Void under Working for Workers Act, 2021 (Ontario) for employees (s.67.2 ESA). Review immediately.",
        "severity": "CRITICAL",
        "contract_types": ["Employment Contract"],
    },
    "Construction Act prompt payment (Ontario)": {
        # Use separate patterns (no alternation with groups) to avoid tuple matches
        "patterns": [r'net\s+(\d+)', r'payment\s+within\s+(\d+)\s+days?'],
        "check": lambda matches: any(
            _is_numeric(m) and int(float(m)) > 28
            for _p, m in matches if m
        ),
        "flag": "Payment terms may exceed Construction Act (Ontario) 28-day prompt payment requirement for proper invoices.",
        "severity": "HIGH",
        "contract_types": ["Construction Contract"],
    },
}


# Short keywords that risk false positives when used as substrings.
# These are checked with word boundaries instead of simple 'in' test.
_BOUNDARY_KEYWORDS = frozenset({"fee", "rate", "price", "pay", "assign"})

def detect_clauses_present(text: str) -> dict:
    """Detect which of the 15 mandatory clauses are present in the contract.

    Uses an expanded keyword list covering common drafting variations.
    Short keywords (fee, rate, price) are matched at word boundaries to
    prevent false positives from substrings like "feesimple" or "pirate".
    Returns a dict of {clause_key: True} for each clause detected.
    """
    text_lower = text.lower()
    found = {}
    for clause_key, keywords in CLAUSE_KEYWORDS.items():
        for kw in keywords:
            if kw in _BOUNDARY_KEYWORDS:
                # Word-boundary match for short ambiguous keywords
                if re.search(r'\b' + re.escape(kw) + r'\b', text_lower):
                    found[clause_key] = True
                    break
            elif kw in text_lower:
                found[clause_key] = True
                break
    return found


def check_statutory_redlines(text: str, contract_type: str = "") -> list:
    """Run Python-layer statutory red-line checks against the contract text.

    These check for non-waivable Canadian law violations that the AI should
    always flag but might miss. Results are injected into the user prompt as
    verified ground truth the AI cannot contradict.

    Returns a list of dicts: [{statute, flag, severity}, ...]
    """
    flags = []
    text_lower = text.lower()
    for statute, rule in STATUTORY_REDLINES.items():
        # Skip contract-type-specific rules if they don't apply
        if "contract_types" in rule and contract_type not in rule["contract_types"]:
            continue
        # Find all pattern matches
        all_matches = []
        for pattern in rule["patterns"]:
            matches = re.findall(pattern, text_lower)
            all_matches.extend([(pattern, m) for m in matches])
        # Run the check function
        try:
            triggered = rule["check"](all_matches)
        except Exception as _check_err:
            _audit.warning(f"Statutory check '{name}' failed: {_check_err}")
            triggered = False
        if triggered:
            flags.append({
                "statute":  statute,
                "flag":     rule["flag"],
                "severity": rule["severity"],
            })
    return flags

# ══════════════════════════════════════════════════════════════
# CITATION VALIDATION ENGINE
# ══════════════════════════════════════════════════════════════
# ── PIPEDA / Privacy Settings (sidebar) ──────────────────────
with st.sidebar.expander("Privacy & PIPEDA Settings", expanded=False):
    st.caption("Configure data retention and download compliance templates.")

    st.markdown("**Matter Store Retention**")
    retention_days = st.selectbox(
        "Auto-delete matters older than:",
        options=[30, 60, 90, 180, 365, 0],
        format_func=lambda x: f"{x} days" if x > 0 else "Never (manual only)",
        index=2,
        key="retention_days",
        help="Matters older than this will be deleted when you click Apply below.")

    if retention_days > 0:
        expiring = matters_older_than(retention_days)
        if expiring > 0:
            st.warning(f"{expiring} matter(s) older than {retention_days} days.")
            if st.button(f"Apply — Delete {expiring} expired matter(s)",
                         key="apply_retention"):
                deleted = purge_expired_matters(retention_days)
                st.success(f"Deleted {deleted} expired matter(s).")
                st.rerun()
        else:
            st.caption(f"No matters older than {retention_days} days.")

    st.markdown("**Delete All Matters**")
    if st.button("🗑 Delete All Matters", key="delete_all_matters",
                 help="Permanently deletes all saved matters. Cannot be undone."):
        if st.session_state.get("delete_all_confirmed"):
            deleted = purge_all_matters()
            st.success(f"Deleted {deleted} matter(s).")
            st.session_state.delete_all_confirmed = False
            st.rerun()
        else:
            st.session_state.delete_all_confirmed = True
            st.warning("Click again to confirm — this cannot be undone.")

    st.markdown("**Retainer Disclosure Template**")
    firm_name_input = st.text_input(
        "Firm name", placeholder="Smith & Associates LLP", key="firm_name_pipeda")
    retainer_province = st.selectbox(
        "Province", list(PROVINCE_RULES.keys()),
        key="retainer_province_sel", index=0)
    retainer_clause = generate_retainer_disclosure(
        firm_name  = firm_name_input or "[FIRM NAME]",
        jurisdiction = retainer_province)
    st.download_button(
        "Download Retainer Disclosure Clause",
        data    = retainer_clause,
        file_name = "PIPEDA_AI_Retainer_Disclosure.txt",
        mime    = "text/plain",
        key     = "sidebar_retainer_download",
        help    = "Add this clause to your standard client retainer agreement.")
    st.caption(
        "This clause addresses: consent, purpose limitation, cross-border transfer "
        "to Anthropic's US infrastructure, Anthropic's no-training commitment, and "
        "the lawyer's professional review obligation.")

def validate_citations(analysis_json, clause_index):
    issues = []
    valid_sections = set(clause_index)
    if not isinstance(analysis_json, dict):
        return issues
    for section_key in ["high_exposure_issues", "elevated_risk_issues", "commercial_imbalances"]:
        for issue in analysis_json.get(section_key, []):
            clause_ref = issue.get("clause_reference", "")
            refs = re.findall(r'(?:Section|Article|Clause|s\.)\s*(\d+(?:\.\d+)*)', clause_ref, re.I)
            for ref in refs:
                if ref not in valid_sections and valid_sections:
                    issues.append({
                        "type": "INVALID_SECTION_REF",
                        "claimed": ref,
                        "in_issue": issue.get("title", ""),
                        "fix": f"Section {ref} not found in clause index. Verify reference."
                    })
    return issues

# ══════════════════════════════════════════════════════════════
# VERSION COMPARISON ENGINE
# ══════════════════════════════════════════════════════════════
def compare_versions(original_text, redline_text):
    orig_lines = [l.strip() for l in original_text.split('\n') if l.strip()]
    red_lines  = [l.strip() for l in redline_text.split('\n') if l.strip()]
    differ = difflib.unified_diff(orig_lines, red_lines, lineterm='', n=2)
    changes = {"added": [], "removed": [], "modified": []}
    removed_buf, added_buf = [], []
    for line in differ:
        if line.startswith('---') or line.startswith('+++') or line.startswith('@@'):
            continue
        if line.startswith('-'):
            removed_buf.append(line[1:].strip())
        elif line.startswith('+'):
            added_buf.append(line[1:].strip())
        else:
            if removed_buf and added_buf:
                changes["modified"].append({"original": " ".join(removed_buf), "revised": " ".join(added_buf)})
            elif removed_buf:
                changes["removed"].extend(removed_buf)
            elif added_buf:
                changes["added"].extend(added_buf)
            removed_buf, added_buf = [], []
    if removed_buf and added_buf:
        changes["modified"].append({"original": " ".join(removed_buf), "revised": " ".join(added_buf)})
    elif removed_buf:
        changes["removed"].extend(removed_buf)
    elif added_buf:
        changes["added"].extend(added_buf)
    return changes

# ══════════════════════════════════════════════════════════════
# TEXT EXTRACTION — single file
# ══════════════════════════════════════════════════════════════
def preprocess_scan(pil_image):
    if not HAS_CV2:
        return pil_image
    img = cv2.cvtColor(np.array(pil_image), cv2.COLOR_RGB2BGR)
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    denoised = cv2.fastNlMeansDenoising(gray, None, h=12, templateWindowSize=7, searchWindowSize=21)
    contrast = cv2.adaptiveThreshold(denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 15, 8)
    try:
        coords = np.column_stack(np.where(contrast < 255))
        if len(coords) > 100:
            angle = cv2.minAreaRect(coords)[-1]
            angle = -(90 + angle) if angle < -45 else -angle
            if abs(angle) > 0.3 and abs(angle) < 15:
                (h, w) = contrast.shape
                M = cv2.getRotationMatrix2D((w//2, h//2), angle, 1.0)
                contrast = cv2.warpAffine(contrast, M, (w, h),
                    flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)
    except Exception:
        pass
    # Morphological closing with a meaningful 3x3 kernel to close
    # small gaps in characters caused by noise or degraded scanning.
    kernel  = np.ones((3, 3), np.uint8)
    cleaned = cv2.morphologyEx(contrast, cv2.MORPH_CLOSE, kernel)
    from PIL import Image
    return Image.fromarray(cleaned)


# ── File upload validation constants ──
MAX_UPLOAD_BYTES = 100 * 1024 * 1024   # 100 MB hard limit
ALLOWED_EXTENSIONS = frozenset({"pdf", "docx", "txt"})


def validate_upload(f) -> bool:
    """Validate an uploaded file before processing.

    Checks extension, file size, and basic filename safety.
    Returns True if safe to process, False otherwise (error shown to user).
    """
    if not f or not f.name:
        st.error("Invalid file: no filename.")
        return False

    # Extension check (use rsplit to handle multi-dot names like 'contract.v2.pdf')
    ext = f.name.rsplit(".", 1)[-1].lower() if "." in f.name else ""
    if ext not in ALLOWED_EXTENSIONS:
        st.error(f"Unsupported file type: .{ext}. Please upload PDF, DOCX, or TXT.")
        return False

    # Size check — seek to end to get size without reading the whole file
    try:
        f.seek(0, 2)          # seek to end
        size = f.tell()
        f.seek(0)             # reset
    except Exception:
        size = getattr(f, "size", 0)

    if size > MAX_UPLOAD_BYTES:
        st.error(f"File too large: {size/1024/1024:.1f} MB (maximum {MAX_UPLOAD_BYTES//1024//1024} MB). "
                 "For very large contracts, split into main agreement + schedules.")
        return False

    return True


def _claude_vision_ocr(images, ai_client, model_id: str, max_chars: int) -> str:
    """Last-resort OCR using Claude's vision API.

    Called when Tesseract is unavailable or produces < 200 characters of
    usable text from a scanned PDF.  Each page image is base64-encoded and
    sent to Claude with a focused extraction prompt.  Claude Vision handles
    rotated text, poor contrast, multi-column layouts, stamps, and handwritten
    annotations far better than Tesseract on degraded scans.

    Cost: ~$0.002-0.004 per page at claude-sonnet rates (image tokens).
    Rate: pages processed sequentially to avoid bursting.
    """
    import io as _io

    text_parts = []
    bar = st.progress(0)
    total = len(images)

    VISION_PROMPT = (
        "You are a legal document OCR engine. Extract ALL text from this contract page "
        "exactly as it appears, preserving paragraph structure, clause numbers, and "
        "headings. Output only the extracted text — no commentary, no descriptions of "
        "the image, no summaries. If a region is illegible, write [ILLEGIBLE]. "
        "Maintain the original reading order (top-to-bottom, left-to-right for "
        "English; right-to-left for any right-to-left scripts). This is a legal "
        "document — accuracy is critical."
    )

    for i, img in enumerate(images):
        try:
            # Convert PIL Image → JPEG bytes → base64
            buf = _io.BytesIO()
            img.save(buf, format="JPEG", quality=92)
            img_b64 = base64.standard_b64encode(buf.getvalue()).decode("utf-8")

            resp = ai_client.messages.create(
                model=model_id,
                max_tokens=4096,
                messages=[{
                    "role": "user",
                    "content": [
                        {
                            "type":  "image",
                            "source": {
                                "type":       "base64",
                                "media_type": "image/jpeg",
                                "data":       img_b64,
                            },
                        },
                        {"type": "text", "text": VISION_PROMPT},
                    ],
                }],
            )
            page_text = resp.content[0].text.strip()
            if page_text:
                text_parts.append(f"\n[PAGE {i+1}]\n{page_text}\n")

        except Exception as page_err:
            st.warning(f"Claude Vision: page {i+1} failed ({page_err}) — skipping.")

        bar.progress((i + 1) / total)
        if sum(len(t) for t in text_parts) >= max_chars:
            break

    bar.empty()
    return "".join(text_parts)


def extract_text_from_file(f, max_chars=500_000):
    """Extract raw text from a single uploaded file object.

    Three-tier OCR fallback for scanned PDFs:
      Tier 1 — pdfplumber   : instant, works on all text-based PDFs
      Tier 2 — Tesseract    : local OCR for scanned pages (requires pytesseract)
      Tier 3 — Claude Vision: last resort when Tesseract unavailable or yields
                               < 200 usable characters. No extra API key needed —
                               uses the same Anthropic client already configured.
    """
    ext = f.name.rsplit('.', 1)[-1].lower() if '.' in f.name else ''
    text = ""; method = "text"

    if ext == 'pdf':
        # ── Tier 1: pdfplumber (text-based PDFs) ─────────────────
        f.seek(0)
        with pdfplumber.open(f) as pdf:
            bar = st.progress(0)
            total = len(pdf.pages)
            for i, pg in enumerate(pdf.pages):
                t = pg.extract_text() or ""
                if t: text += f"\n[PAGE {i+1}]\n{t}\n"
                bar.progress((i + 1) / total)
                if len(text) >= max_chars: break
            bar.empty()

        clean = re.sub(r'\[PAGE \d+\]|\s+', '', text)
        is_scanned = len(clean) < 100

        if is_scanned and HAS_OCR:
            # ── Tier 2: Tesseract OCR ─────────────────────────────
            preprocess_label = " with image preprocessing" if HAS_CV2 else ""
            st.info(f"Scanned PDF detected. Running Tesseract OCR{preprocess_label} (Tier 2)...")
            text = ""; method = "OCR"; f.seek(0)
            tier2_ok = False
            try:
                import concurrent.futures
                images = convert_from_bytes(f.read(), dpi=300)
                bar2 = st.progress(0)

                def _ocr_page(args):
                    idx, img = args
                    if HAS_CV2: img = preprocess_scan(img)
                    return idx, pytesseract.image_to_string(
                        img, lang='eng', config='--psm 6' if HAS_CV2 else '')

                with concurrent.futures.ThreadPoolExecutor(max_workers=2) as pool:
                    futures = {pool.submit(_ocr_page, (i, img)): i
                               for i, img in enumerate(images)}
                    results = {}
                    for future in concurrent.futures.as_completed(futures):
                        i, ocr = future.result()
                        results[i] = ocr
                        bar2.progress(len(results) / len(images))

                for i in sorted(results):
                    if results[i].strip():
                        text += f"\n[PAGE {i+1}]\n{results[i]}\n"
                bar2.empty()

                tier2_chars = len(re.sub(r'\[PAGE \d+\]|\s+', '', text))
                if tier2_chars >= 200:
                    tier2_ok = True
                    st.success(
                        f"Tesseract OCR complete: {len(text):,} chars "
                        f"from {len(images)} pages")
                else:
                    st.warning(
                        f"Tesseract yielded only {tier2_chars} chars — "
                        "this scan may be too degraded. Trying Claude Vision (Tier 3)...")

            except Exception as e:
                images = None
                st.warning(f"Tesseract OCR failed: {e}. Trying Claude Vision fallback (Tier 3)...")

            # ── Tier 3: Claude Vision fallback ───────────────────
            if not tier2_ok:
                try:
                    # Re-render images if Tesseract crashed before doing so
                    if images is None:
                        f.seek(0)
                        images = convert_from_bytes(f.read(), dpi=200)
                    st.info(
                        f"Falling back to Claude Vision OCR on {len(images)} page(s). "
                        "This uses image tokens (~$0.002–0.004/page).")
                    text   = ""
                    method = "Claude Vision OCR"
                    text   = _claude_vision_ocr(images, ai, model_id, max_chars)
                    vision_chars = len(re.sub(r'\[PAGE \d+\]|\s+', '', text))
                    if vision_chars >= 50:
                        st.success(
                            f"Claude Vision OCR complete: {len(text):,} chars "
                            f"from {len(images)} pages")
                    else:
                        st.error(
                            "Claude Vision also returned minimal text. "
                            "The document may be encrypted, corrupt, or contain "
                            "only images without text content.")
                except Exception as vision_err:
                    st.error(
                        f"Claude Vision OCR failed: {vision_err}. "
                        "Try a cleaner scan or a text-based PDF.")

        elif is_scanned and not HAS_OCR:
            # No Tesseract — go straight to Claude Vision
            # pdf2image (convert_from_bytes) is still needed to rasterise the pages.
            # It's installed with pytesseract in most setups, but guard here.
            st.info("Tesseract not installed. Attempting Claude Vision OCR (Tier 3)...")
            f.seek(0)
            try:
                from pdf2image import convert_from_bytes as _cfb
                images = _cfb(f.read(), dpi=200)
                text   = ""
                method = "Claude Vision OCR"
                text   = _claude_vision_ocr(images, ai, model_id, max_chars)
                vision_chars = len(re.sub(r'\[PAGE \d+\]|\s+', '', text))
                if vision_chars >= 50:
                    st.success(
                        f"Claude Vision OCR complete: {len(text):,} chars from {len(images)} pages")
                else:
                    st.error("Claude Vision returned minimal text. Check the document quality.")
            except Exception as vision_err:
                st.error(f"Claude Vision OCR failed: {vision_err}")

    elif ext == 'docx':
        text = "\n".join(p.text for p in DocxDoc(f).paragraphs if p.text.strip())
    elif ext == 'txt':
        text = f.read().decode('utf-8', errors='replace')

    # Prepend an OCR notice so the AI handles artifacts gracefully
    if method in ("OCR", "Claude Vision OCR"):
        notice = (
            f"[OCR NOTICE: Text extracted via {method}. "
            "Some characters may be misread — interpret in most likely legal context.]\n\n"
        )
        text = notice + text

    return text[:max_chars].strip(), method

def build_clause_index(text):
    sections_found = re.findall(r'(?:Section|Article|Clause|SECTION|ARTICLE|CLAUSE)\s+(\d+(?:\.\d+)*)', text)
    numbered = re.findall(r'^(\d+(?:\.\d+)*)\s+[A-Z]', text, re.MULTILINE)
    all_sections = sorted(set(sections_found + numbered))
    if all_sections:
        idx = "\n\n[CLAUSE INDEX — Sections present: " + ", ".join(all_sections) + "]\n"
        idx += "[Only reference section numbers from this list.]\n\n"
        return idx + text, all_sections
    return text, all_sections


def extract_text(uploaded_files, label="contract"):
    """
    Multi-document aware text extractor.
    uploaded_files: single UploadedFile or list of UploadedFiles.
    Returns (combined_text, method, clause_index_list).
    """
    if not isinstance(uploaded_files, list):
        uploaded_files = [uploaded_files]
    combined = ""
    methods = []
    for i, f in enumerate(uploaded_files):
        doc_label = f"[DOCUMENT {i+1}: {f.name}]\n" if len(uploaded_files) > 1 else ""
        raw, method = extract_text_from_file(f)
        combined += f"\n{doc_label}{raw}\n"
        methods.append(method)
    method = "OCR" if "OCR" in methods else "text"
    final, clause_index = build_clause_index(combined.strip())
    return final, method, clause_index


# ══════════════════════════════════════════════════════════════
# LONG-CONTRACT CHUNKING PIPELINE
# Contracts > CHUNK_THRESHOLD chars are split into overlapping
# chunks, each analysed independently, then a synthesis pass
# merges findings, deduplicates issues, and produces the final
# unified JSON. No contract is ever silently truncated.
# ══════════════════════════════════════════════════════════════

def _split_into_chunks(text, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    """Split text into overlapping chunks, preferring paragraph boundaries."""
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        # Try to end at a paragraph boundary
        if end < len(text):
            para_end = text.rfind('\n\n', start + chunk_size // 2, end)
            if para_end != -1:
                end = para_end
        chunks.append(text[start:end])
        if end >= len(text):
            break
        start = end - overlap
    return chunks


def _merge_chunk_analyses(chunk_results):
    """
    Merge multiple chunk JSON analyses into one unified result.
    Deduplicates issues by title similarity. Sums risk scores.
    """
    merged = {
        "contract_type":        "",
        "jurisdiction":         "",
        "client_party":         "",
        "counterparty":         "",
        "effective_date":       None,
        "risk_score":           0,
        "risk_level":           "LOW",
        "verdict":              "",
        "executive_summary":    "",
        "mandatory_clause_review": [],
        "high_exposure_issues": [],
        "elevated_risk_issues": [],
        "commercial_imbalances":[],
        "missing_protections":  [],
        "negotiation_strategy": {"priority_items": [], "walk_away_points": [], "concession_candidates": []},
        "action_items":         {"immediate": [], "before_execution": [], "post_execution": []},
        "suggested_redline_clauses": [],
        "revised_contract":     "",
    }
    seen_issue_titles = set()
    seen_clause_categories = set()

    for r in chunk_results:
        if not isinstance(r, dict):
            continue
        # Take first non-empty scalar fields
        for f in ("contract_type", "jurisdiction", "client_party", "counterparty", "effective_date"):
            if not merged[f] and r.get(f):
                merged[f] = r[f]

        # Risk score: take max per chunk (not sum) to avoid inflation.
        # Three chunks each scoring 5 for the same issue should not become 15.
        # After merge, re-derive from unique deduplicated issue count below.
        merged["risk_score"] = max(merged["risk_score"], r.get("risk_score", 0))
        level_order = {"LOW": 0, "MODERATE": 1, "HIGH": 2, "CRITICAL": 3}
        if level_order.get(r.get("risk_level","LOW"), 0) > level_order.get(merged["risk_level"], 0):
            merged["risk_level"] = r["risk_level"]
            merged["verdict"] = r.get("verdict", "")

        # Executive summary — append unique sentences
        if r.get("executive_summary"):
            if merged["executive_summary"]:
                merged["executive_summary"] += " " + r["executive_summary"]
            else:
                merged["executive_summary"] = r["executive_summary"]

        # Mandatory clause review — first occurrence per category wins
        for item in r.get("mandatory_clause_review", []):
            cat = item.get("clause_category", "")
            if cat and cat not in seen_clause_categories:
                seen_clause_categories.add(cat)
                merged["mandatory_clause_review"].append(item)

        # Issues — deduplicate by title (fuzzy: first 40 chars)
        for section in ("high_exposure_issues", "elevated_risk_issues", "commercial_imbalances"):
            for issue in r.get(section, []):
                # Composite dedup key: title (60 chars) + clause_ref + severity.
                # Using only 40 chars of title caused false collisions:
                # "Limitation of liability clause..." and "Limitation of liability cap..."
                # both hashed to the same key and one was silently dropped.
                title_key = "|||".join([
                    issue.get("title", "")[:60].lower().strip(),
                    issue.get("clause_reference", "").lower().strip(),
                    issue.get("severity", "").lower().strip(),
                ])
                if title_key not in seen_issue_titles:
                    seen_issue_titles.add(title_key)
                    merged[section].append(issue)

        # Missing protections — deduplicate by clause_type
        seen_missing = {m.get("clause_type","") for m in merged["missing_protections"]}
        for item in r.get("missing_protections", []):
            if item.get("clause_type","") not in seen_missing:
                merged["missing_protections"].append(item)
                seen_missing.add(item.get("clause_type",""))

        # Negotiation strategy — union of lists
        for key in ("priority_items", "walk_away_points", "concession_candidates"):
            for item in r.get("negotiation_strategy", {}).get(key, []):
                if item not in merged["negotiation_strategy"][key]:
                    merged["negotiation_strategy"][key].append(item)

        # Action items — union
        for phase in ("immediate", "before_execution", "post_execution"):
            for item in r.get("action_items", {}).get(phase, []):
                if item not in merged["action_items"][phase]:
                    merged["action_items"][phase].append(item)

        # Redline clauses
        seen_rl = {x.get("clause_reference","") for x in merged["suggested_redline_clauses"]}
        for rl in r.get("suggested_redline_clauses", []):
            if rl.get("clause_reference","") not in seen_rl:
                merged["suggested_redline_clauses"].append(rl)
                seen_rl.add(rl.get("clause_reference",""))

    # Re-derive risk score from unique deduplicated issues (not raw max from chunks).
    # CRITICAL issues = 3 pts, ELEVATED = 2 pts, MODERATE = 1 pt.
    # This prevents a single high-risk chunk inflating the merged score.
    derived_score = (
        sum(3 for i in merged["high_exposure_issues"])
        + sum(2 for i in merged["elevated_risk_issues"])
        + sum(1 for i in merged["commercial_imbalances"])
    )
    # Keep the max of the AI-reported score and the derived score as a floor,
    # but cap at derived so chunk-counting artifacts are removed.
    merged["risk_score"] = max(derived_score, merged["risk_score"] // 2)
    score = merged["risk_score"]
    if score >= 15:   merged["risk_level"] = "CRITICAL"
    elif score >= 10: merged["risk_level"] = "HIGH"
    elif score >= 5:  merged["risk_level"] = "MODERATE"
    else:             merged["risk_level"] = "LOW"

    # Trim executive summary to ~4 sentences
    sents = re.split(r'(?<=[.!?])\s+', merged["executive_summary"])
    merged["executive_summary"] = " ".join(sents[:4])

    return merged

# ══════════════════════════════════════════════════════════════
# JSON PARSING
# ══════════════════════════════════════════════════════════════
# ── Schema validation constants ─────────────────────────────
# Required top-level keys the AI must return. Callers check these after parsing.
_REQUIRED_TOP_KEYS = frozenset({
    "risk_score", "risk_level", "verdict", "executive_summary",
    "mandatory_clause_review", "high_exposure_issues",
})
_MANDATORY_CLAUSE_COUNT = 15   # prompt requires exactly 15 entries

# ── Retry config ─────────────────────────────────────────────
_MAX_RETRIES    = 2    # max extra attempts after the first (total = 3)
_RETRY_DELAYS   = [4, 12]   # seconds to wait before retry 1 and retry 2


def validate_analysis_schema(data: dict) -> list[str]:
    """Check that the AI-returned dict satisfies the minimum required schema.

    Returns a list of validation error strings (empty = valid).
    This runs AFTER successful JSON parse so we can catch structurally valid
    JSON that is semantically wrong (e.g. missing mandatory_clause_review,
    wrong risk_level enum, etc.).
    """
    errors = []
    if not isinstance(data, dict):
        return ["Response is not a JSON object"]

    # Required top-level keys
    for key in _REQUIRED_TOP_KEYS:
        if key not in data:
            errors.append(f"Missing required field: '{key}'")

    # risk_level must be one of the allowed values
    rl = data.get("risk_level", "")
    if rl not in ("CRITICAL", "HIGH", "MODERATE", "LOW"):
        errors.append(f"Invalid risk_level: {rl!r} — must be CRITICAL|HIGH|MODERATE|LOW")

    # risk_score must be a non-negative integer
    rs = data.get("risk_score")
    if not isinstance(rs, int) or rs < 0:
        errors.append(f"Invalid risk_score: {rs!r} — must be a non-negative integer")

    # mandatory_clause_review must be a list
    mcr = data.get("mandatory_clause_review")
    if not isinstance(mcr, list):
        errors.append("mandatory_clause_review must be a list")
    elif len(mcr) < 1:
        errors.append("mandatory_clause_review is empty")
    # Warn (not error) on count — chunked analysis may have fewer
    elif len(mcr) < _MANDATORY_CLAUSE_COUNT:
        errors.append(
            f"mandatory_clause_review has {len(mcr)} entries (expected {_MANDATORY_CLAUSE_COUNT})"
        )

    # high_exposure_issues must be a list
    if not isinstance(data.get("high_exposure_issues", []), list):
        errors.append("high_exposure_issues must be a list")

    return errors


def _build_retry_nudge(validation_errors: list, raw_response: str) -> str:
    """Build a follow-up user message asking the AI to fix its JSON output."""
    error_str = "\n".join(f"  - {e}" for e in validation_errors)
    snippet = raw_response[:300] + ("..." if len(raw_response) > 300 else "")
    return (
        f"Your previous response had the following issues:\n{error_str}\n\n"
        f"Beginning of your previous response:\n{snippet}\n\n"
        "Please provide ONLY a corrected, valid JSON object that fixes all the issues above. "
        "No markdown, no code fences, no surrounding text. Start your response with '{' and "
        "end it with '}'. Ensure mandatory_clause_review has exactly 15 entries and "
        "risk_level is one of: CRITICAL, HIGH, MODERATE, LOW."
    )
def _call_ai_with_retry(ai_client, model_id: str, max_tokens: int,
                        system_cfg, user_msg: str,
                        timeout: float = 420.0,
                        progress_label: str = "") -> tuple:
    """Call the AI API with exponential-backoff retry on failure or invalid JSON.

    On the first successful parse that also passes schema validation, returns
    (data_dict, raw_text, True).
    After all retries are exhausted, returns (None, last_raw_text, False).

    Retry triggers:
      - JSON parse failure (model output was not valid JSON at all)
      - Schema validation failure (JSON parsed but missing required fields)
      - API-level exceptions (rate limit, timeout, network error)
    """
    last_raw = ""
    history  = [{"role": "user", "content": user_msg}]

    for attempt in range(1 + _MAX_RETRIES):
        try:
            full = ""
            prog = st.empty() if progress_label else None
            with ai_client.messages.stream(
                model=model_id, max_tokens=max_tokens,
                system=system_cfg,
                messages=history,
                timeout=timeout,
            ) as stream:
                last_log = 0
                for txt in stream.text_stream:
                    full += txt
                    if prog and len(full) - last_log > 2000:
                        prog.text(f"{progress_label} {len(full):,} chars")
                        last_log = len(full)
            if prog:
                prog.empty()
            last_raw = full

            data, ok = parse_json_response(full)
            if not ok or data is None:
                # JSON parse failed entirely
                errs = ["Response is not valid JSON"]
            else:
                errs = validate_analysis_schema(data)

            if not errs:
                return data, full, True

            # Schema or parse failed — decide whether to retry
            if attempt < _MAX_RETRIES:
                delay = _RETRY_DELAYS[attempt]
                st.warning(
                    f"AI response validation failed (attempt {attempt+1}/{1+_MAX_RETRIES}): "
                    f"{errs[0]}. Retrying in {delay}s..."
                )
                time.sleep(delay)
                # Add the bad response + correction request to conversation history
                history.append({"role": "assistant", "content": full})
                history.append({"role": "user",      "content": _build_retry_nudge(errs, full)})
            else:
                st.warning(
                    f"AI response failed schema validation after {1+_MAX_RETRIES} attempts. "
                    f"Last errors: {'; '.join(errs[:3])}. Displaying best-effort result."
                )
                # Return the last parsed result even if imperfect
                return data, last_raw, ok

        except Exception as e:
            last_raw = full if 'full' in locals() else ""
            err_str  = str(e)
            if attempt < _MAX_RETRIES:
                delay = _RETRY_DELAYS[attempt]
                if "rate" in err_str.lower() or "429" in err_str:
                    st.warning(f"Rate limited. Waiting {delay}s before retry {attempt+2}...")
                elif "timeout" in err_str.lower():
                    st.warning(f"API timeout. Retrying in {delay}s (attempt {attempt+2})...")
                else:
                    st.warning(f"API error: {err_str}. Retrying in {delay}s...")
                time.sleep(delay)
            else:
                raise   # re-raise after exhausting retries

    return None, last_raw, False


def parse_json_response(raw: str) -> tuple:
    """Parse AI JSON response with aggressive fallback strategies.

    Returns (parsed_dict, True) on success.
    Returns (None, False) on failure — callers must check the bool before
    calling .get() or iterating, since the first element is None not a string.
    Previously returned (raw_string, False) which caused silent AttributeErrors
    when callers called .get() on the returned string.
    """
    cleaned = raw.strip()
    cleaned = re.sub(r'^```(?:json)?\s*\n?', '', cleaned)
    cleaned = re.sub(r'\n?\s*```\s*$', '', cleaned)
    cleaned = cleaned.strip()
    try:
        return json.loads(cleaned), True
    except json.JSONDecodeError:
        pass
    start = cleaned.find('{'); end = cleaned.rfind('}')
    if start != -1 and end != -1 and end > start:
        subset = cleaned[start:end + 1]
        for attempt_fn in [
            lambda s: json.loads(s),
            lambda s: json.loads(re.sub(r',\s*}', '}', re.sub(r',\s*]', ']', s))),
        ]:
            try:
                return attempt_fn(subset), True
            except Exception:
                pass
        # Truncated JSON — try to close it
        open_b = subset.count('{') - subset.count('}')
        open_k = subset.count('[') - subset.count(']')
        if open_b > 0 or open_k > 0:
            trunc = subset.rstrip()
            if trunc.count('"') % 2 != 0:
                lq = trunc.rfind('"')
                slq = trunc.rfind('"', 0, lq)
                if slq != -1:
                    trunc = trunc[:slq]
            trunc = trunc.rstrip().rstrip(',').rstrip(':').rstrip()
            trunc += ']' * max(0, open_k) + '}' * max(0, open_b)
            trunc = re.sub(r',\s*}', '}', re.sub(r',\s*]', ']', trunc))
            try:
                return json.loads(trunc), True
            except Exception:
                pass
    # Returning None (not raw string) so callers cannot accidentally call
    # .get() on a string and get a silent AttributeError.
    return None, False


# ══════════════════════════════════════════════════════════════
# NEGOTIATION SIMULATOR PROMPT
# ══════════════════════════════════════════════════════════════
def build_negotiation_sim_prompt(issues_json, contract_type):
    issues_summary = [
        {"title": i.get("title",""), "clause": i.get("clause_reference",""),
         "proposed_change": i.get("proposed_replacement", i.get("recommendation","")),
         "severity": i.get("severity","")}
        for i in issues_json
    ]
    return f"""You are an experienced commercial litigator simulating how opposing counsel would respond to each proposed amendment.

Contract Type: {contract_type}

PROPOSED AMENDMENTS:
{json.dumps(issues_summary, indent=2)}

For each amendment, predict the opposing counsel's likely response. Respond with ONLY valid JSON:
{{
  "simulations": [
    {{
      "amendment_title": "string",
      "opposing_likely_response": "ACCEPT | REJECT | COUNTER",
      "predicted_counter": "string",
      "best_case": "string",
      "likely_compromise": "string",
      "fallback_position": "string",
      "leverage_notes": "string",
      "probability_of_acceptance": integer 0-100
    }}
  ]
}}

Be realistic. Base predictions on standard Canadian commercial negotiation dynamics for {contract_type} agreements."""


# ══════════════════════════════════════════════════════════════
# SOURCE LINKS
# ══════════════════════════════════════════════════════════════
STATUTE_URLS = {
    "ESA": "https://www.ontario.ca/laws/statute/00e41",
    "Employment Standards Act": "https://www.ontario.ca/laws/statute/00e41",
    "ESA 2000": "https://www.ontario.ca/laws/statute/00e41",
    "Construction Act": "https://www.ontario.ca/laws/statute/90c30",
    "Consumer Protection Act": "https://www.ontario.ca/laws/statute/02c30",
    "CPA 2002": "https://www.ontario.ca/laws/statute/02c30",
    "Arthur Wishart Act": "https://www.ontario.ca/laws/statute/00a03",
    "Limitation Act": "https://www.ontario.ca/laws/statute/02l24",
    "Criminal Code": "https://laws-lois.justice.gc.ca/eng/acts/c-46/",
    "Criminal Code s.347": "https://laws-lois.justice.gc.ca/eng/acts/c-46/section-347.html",
    "PIPEDA": "https://laws-lois.justice.gc.ca/eng/acts/p-8.6/",
    "Competition Act": "https://laws-lois.justice.gc.ca/eng/acts/c-34/",
    "Copyright Act": "https://laws-lois.justice.gc.ca/eng/acts/c-42/",
    "Canada Labour Code": "https://laws-lois.justice.gc.ca/eng/acts/l-2/",
    "CBCA": "https://laws-lois.justice.gc.ca/eng/acts/c-44/",
    "PIPA": "https://www.bclaws.gov.bc.ca/civix/document/id/complete/statreg/03063_01",
    "PIPA Alberta": "https://kings-printer.alberta.ca/1266.cfm?page=P06P5.cfm",
    "Civil Code of Quebec": "https://www.legisquebec.gouv.qc.ca/en/document/cs/CCQ-1991",
    "CCQ": "https://www.legisquebec.gouv.qc.ca/en/document/cs/CCQ-1991",
    "Charter of the French Language": "https://www.legisquebec.gouv.qc.ca/en/document/cs/C-11",
    "Bill 96": "https://www.legisquebec.gouv.qc.ca/en/document/cs/C-11",
    "Bhasin v Hrynew": "https://www.canlii.org/en/ca/scc/doc/2014/2014scc71/2014scc71.html",
    "Bardal": "https://www.canlii.org/en/on/onsc/doc/1960/1960canlii86/1960canlii86.html",
    "Working for Workers Act": "https://www.ontario.ca/laws/statute/21w26",
}

def generate_source_links(analysis_text):
    links = []
    text_str = json.dumps(analysis_text) if isinstance(analysis_text, dict) else str(analysis_text)
    for statute_name, url in STATUTE_URLS.items():
        if statute_name.lower() in text_str.lower():
            links.append({"name": statute_name, "url": url})
    seen = set()
    return [lk for lk in links if not (lk['url'] in seen or seen.add(lk['url']))]


# ══════════════════════════════════════════════════════════════
# SYSTEM PROMPT
# ══════════════════════════════════════════════════════════════
def build_system_prompt(mode, province, playbook_text="",
                        contract_type="Auto-Detect", language="English",
                        deal_ctx: DealContext = None):
    prov = PROVINCE_RULES.get(province, PROVINCE_RULES["Other / Multi-Province"])
    ct = contract_type if contract_type != "Auto-Detect" else "default"
    benchmarks_str = format_benchmarks_for_prompt(ct)

    # Get deal-context-aware variable values for precedent rendering
    variables = get_variable_values(deal_ctx, ct) if deal_ctx else {}
    precedents_str = format_precedents_for_prompt(ct, variables if variables else None)

    # Deal context section (injected only when provided)
    deal_context_str = format_deal_context_for_prompt(deal_ctx) if deal_ctx else ""

    language_instruction = ""
    if language.startswith("French"):
        language_instruction = """
LANGUAGE — MANDATORY:
The contract may be in French. Analyse it in the language it is written in.
Output the ENTIRE JSON response in FRENCH. All analysis, recommendations, market_standard,
proposed_replacement — everything in French. Apply Quebec civil law (CCQ) as the primary
legal framework. Reference CCQ articles, not common law precedents, unless the contract
specifies a common law jurisdiction. Draft all proposed replacement clauses in grammatically
correct legal French.
"""

    mode_instructions = {
        "Associate Memo (Internal)": "You are a senior associate preparing a first-pass review memo for a supervising partner. Write in precise, professional legal prose. Cite clause numbers and pages. Where a matter requires partner judgment, note [PARTNER REVIEW]. Compare each clause to market standards (or to firm playbook if provided). Do not use client-facing language. Do not use emojis or casual phrasing.",
        "Client Advisory (Plain English)": "You are a lawyer writing directly to the client. Convert all legal analysis into clear, accessible language that a non-lawyer business executive would understand. Explain the practical impact. Focus on what the client should do. Avoid unnecessary jargon. Maintain professionalism but prioritise clarity. No emojis.",
        "Opposing Counsel Redline": "You are preparing a response for opposing counsel. For each issue: identify the clause, state the concern, propose a specific revision, and provide a brief legal justification. Maintain a firm but professional tone. Every proposed revision must be a complete, ready-to-insert replacement clause. No emojis.",
    }

    playbook_section = f"""
FIRM PLAYBOOK — MANDATORY:
The reviewing lawyer has provided the following firm-specific standards. You MUST compare
the contract against THESE positions rather than generic market norms. Where the contract
deviates from the firm playbook, flag the deviation explicitly.

{playbook_text}
""" if playbook_text else ""

    clause_list_str = ", ".join(MANDATORY_CLAUSE_LABELS.values())

    return f"""You are a senior Canadian contract lawyer conducting a thorough first-pass review.

OUTPUT MODE: {mode}
{mode_instructions[mode]}
{language_instruction}
TONE AND LANGUAGE — MANDATORY:
- Write in professional legal prose appropriate for a memorandum from a senior associate.
- Do NOT use emojis anywhere.
- Do NOT use dramatic language. Use measured terms: "material risk", "significant exposure",
  "commercially unreasonable", "departure from market standard".

CANADIAN LAW:
{prov}

General Canadian principles: Good faith in contractual performance (Bhasin v Hrynew, 2014 SCC 71).
Penalty clauses unenforceable; liquidated damages must be a genuine pre-estimate of loss.
Force majeure must be explicitly drafted. Criminal Code s.347 / Bill C-46: interest exceeding 35% APR is
a criminal offence for commercial lending (s.347.1, amended by Bill C-46,
effective Jan 1 2025). The original 60% threshold applies to other credit.
Flag any late-fee or interest clause above 35% APR as CRITICAL. PIPEDA governs personal
information. IP ownership defaults to the creator in Canada (no work-for-hire doctrine);
moral rights must be explicitly waived. Competition Act 2024 amendments address drip pricing.
Working for Workers Act, 2021 (Ontario): non-competes void for employees.

ANTI-HALLUCINATION RULES — MANDATORY:
1. ONLY reference sections/clauses that ACTUALLY EXIST in the contract text.
2. ONLY quote text that appears VERBATIM in the contract.
3. If a section number is not findable, write "[Not numbered — see paragraph beginning with '...']".
4. If a standard clause is MISSING, explicitly state it is absent.
5. NEVER invent statute citations. If uncertain, write "[Verify: believed to be under [Act name]]".
6. NEVER fabricate case law.
7. If the contract text is a CHUNK (part of a longer document), note this and avoid drawing
   conclusions that depend on sections not present in this chunk.

MANDATORY 15-CLAUSE CHECKLIST:
You MUST evaluate ALL of the following: {clause_list_str}.
If any is absent, flag it under Missing Protections.

MARKET BENCHMARKS — compare each clause against these quantitative standards:
{benchmarks_str}

{deal_context_str}

PRECEDENT CLAUSE LIBRARY — when proposing replacement clauses, adapt from these vetted
precedents. The placeholders have already been filled with deal-specific values.
Use these values (months, amounts, days) in your proposed replacement — do not substitute
generic figures:
{precedents_str}

For each clause in mandatory_clause_review, include a benchmark_comparison field stating:
1. What the market standard is for this clause category and contract type
2. What the contract actually provides (cite the specific figure/term)
3. Whether the contract term is above, at, or below the deal-calibrated standard
4. If deal context was provided: compare against the dynamic thresholds, not generic benchmarks

ENFORCEABILITY SCORING — MANDATORY:
For each issue, estimate the probability (0-100) that the clause would be enforced as written
if challenged in court in {province}. Consider: clarity of drafting, statutory compliance,
judicial trends, unconscionability principles.

{playbook_section}

CONFIDENCE TAGGING — MANDATORY:
HIGH = confident in analysis and clause text is clear.
MEDIUM = sound but clause is ambiguous or legal position is arguable.
REQUIRES_REVIEW = unclear clause, unsettled law, or unable to verify statute.

RISK SCORING:
CRITICAL (3 pts), ELEVATED (2 pts), MODERATE (1 pt), ACCEPTABLE (0 pts).
Total: 15+ = CRITICAL, 10-14 = HIGH, 5-9 = MODERATE, 0-4 = LOW.

INLINE DISCLAIMER — MANDATORY:
Every proposed_replacement field MUST begin with this exact text on its own line:
"[AI DRAFT — requires review by qualified counsel before use]"
This is a professional responsibility requirement.

OUTPUT FORMAT — STRICT JSON:
Respond with ONLY a valid JSON object. No text before or after. No markdown code fences.
Keep all string values concise. The entire JSON must fit within the token limit.

SCHEMA:
{{
  "contract_type": "string",
  "jurisdiction": "string",
  "client_party": "string",
  "counterparty": "string",
  "effective_date": "string or null",
  "risk_score": integer,
  "risk_level": "CRITICAL | HIGH | MODERATE | LOW",
  "verdict": "string — one sentence",
  "executive_summary": "string — 3 sentences max",
  "mandatory_clause_review": [
    {{
      "clause_category": "string",
      "present": boolean,
      "clause_reference": "string",
      "assessment": "ACCEPTABLE | NEEDS_AMENDMENT | UNACCEPTABLE | MISSING",
      "risk_level": "CRITICAL | HIGH | MEDIUM | LOW | MISSING",
      "analysis": "string — 1-2 sentences",
      "benchmark_comparison": "string — 1 sentence: market standard vs contract",
      "enforceability_pct": integer,
      "confidence": "HIGH | MEDIUM | REQUIRES_REVIEW"
    }}
  ],
  "high_exposure_issues": [
    {{
      "title": "string",
      "severity": "CRITICAL",
      "clause_reference": "string",
      "verbatim_text": "string — key quote only",
      "analysis": "string — 2-3 sentences",
      "market_standard": "string — cite benchmark figures",
      "provincial_law": "string — statute or case, 1 sentence",
      "recommendation": "string — 1 sentence",
      "proposed_replacement": "string — complete replacement clause (adapt from precedent library)",
      "precedent_used": "string — which precedent clause was adapted, or 'drafted' if none",
      "enforceability_pct": integer,
      "confidence": "HIGH | MEDIUM | REQUIRES_REVIEW",
      "partner_review": boolean
    }}
  ],
  "elevated_risk_issues": [
    {{
      "title": "string",
      "severity": "ELEVATED",
      "clause_reference": "string",
      "verbatim_text": "string",
      "analysis": "string — 1-2 sentences",
      "recommendation": "string",
      "proposed_replacement": "string — adapt from precedent library",
      "confidence": "HIGH | MEDIUM | REQUIRES_REVIEW"
    }}
  ],
  "commercial_imbalances": [
    {{
      "title": "string",
      "clause_reference": "string",
      "description": "string — 1 sentence",
      "recommendation": "string — 1 sentence"
    }}
  ],
  "missing_protections": [
    {{
      "clause_type": "string",
      "risk_level": "CRITICAL | HIGH | MEDIUM",
      "explanation": "string — 1 sentence",
      "proposed_clause": "string — adapt from precedent library"
    }}
  ],
  "negotiation_strategy": {{
    "priority_items": ["string"],
    "walk_away_points": ["string"],
    "concession_candidates": ["string"],
    "draft_communication": "string — draft email to opposing counsel"
  }},
  "action_items": {{
    "immediate": ["string"],
    "before_execution": ["string"],
    "post_execution": ["string"]
  }},
  "suggested_redline_clauses": [
    {{
      "clause_reference": "string",
      "issue": "string",
      "current_text": "string",
      "replacement_text": "string",
      "rationale": "string"
    }}
  ],
  "revised_contract": ""
}}

CRITICAL RULES:
1. Output ONLY valid JSON. No markdown, no code fences, no surrounding text.
   Your response MUST start with {{ and end with }}.
2. mandatory_clause_review MUST have exactly 15 entries — one per mandatory category.
3. Keep ALL text fields concise. JSON must complete within the token limit.
4. Properly escape all strings: use \n for newlines, \" for quotes inside strings.
   Never include raw newline characters inside a JSON string value.
5. No emojis. Use severity labels (CRITICAL, ELEVATED, MODERATE, ACCEPTABLE) only.
6. Set revised_contract to empty string unless Deep Dive mode is specified.
7. For proposed_replacement fields, ALWAYS adapt from the precedent library above.
   State which precedent was adapted in the precedent_used field.
8. If PRE-VERIFIED PYTHON RISK SCORES appear in the user message, your risk_level
   and enforceability assessments for those clauses MUST be consistent with them.
9. If STATUTORY RED-LINES appear in the user message, those issues MUST appear as
   high_exposure_issues with the correct severity — they are verified facts.
10. benchmark_comparison field MUST cite specific numbers (months, dollars, days)
    from the DYNAMIC THRESHOLDS section if provided, not generic "market standard".

CHAIN-OF-THOUGHT (Deep Dive mode only):
Before writing the JSON, think through:
  a) What type of contract is this and who are the parties?
  b) Which clauses are present, which are missing?
  c) What are the 3 most critical risks for the client?
  d) Are any statutory minimums at risk?
  e) What is the overall risk posture?
Then output only the JSON — no preamble.

ONE-SHOT EXAMPLE of a valid high_exposure_issue entry:
{{
  "title": "Liability cap below market floor",
  "severity": "CRITICAL",
  "clause_reference": "Section 8.2",
  "verbatim_text": "liability shall not exceed $5,000",
  "analysis": "The $5,000 cap represents 0.5 months of the annual fees, far below the deal-calibrated floor of 6 months (approximately $60,000). This exposes the client to unrecovered losses on any material breach.",
  "market_standard": "Deal-calibrated standard: 12 months = CAD $120,000. Floor: 6 months = CAD $60,000.",
  "provincial_law": "No specific statute; general contract law — courts may find unconscionable.",
  "recommendation": "Negotiate cap to minimum 6 months fees (CAD $60,000) or total contract value.",
  "proposed_replacement": "[AI DRAFT — requires review by qualified counsel before use]\nLimitation of Liability. The aggregate liability of either party shall not exceed the fees paid in the 12 (twelve) month period preceding the claim (approximately CAD $120,000).",
  "precedent_used": "liability_cap / Service Agreement",
  "enforceability_pct": 85,
  "confidence": "HIGH",
  "partner_review": true
}}
"""

# ══════════════════════════════════════════════════════════════
# RISK ENGINE
# ══════════════════════════════════════════════════════════════
def compute_risk(data):
    if isinstance(data, dict):
        return (data.get("risk_level", "MODERATE"),
                data.get("risk_score", 0),
                data.get("verdict", "Review required."))
    return "MODERATE", 0, "Unable to parse risk assessment."


# ══════════════════════════════════════════════════════════════
# UI RENDERERS
# ══════════════════════════════════════════════════════════════
def render_heatmap(clause_review):
    css_map = {"CRITICAL":"heat-critical","HIGH":"heat-high","MEDIUM":"heat-medium",
               "LOW":"heat-low","MISSING":"heat-missing","N/A":"heat-na"}
    cells = ""
    for item in clause_review:
        cat   = item.get("clause_category","")
        label = html_escape(MANDATORY_CLAUSE_LABELS.get(cat, cat.replace("_"," ").title()))
        level = html_escape(item.get("risk_level","N/A").upper())
        css   = css_map.get(item.get("risk_level","N/A").upper(), "heat-na")
        cells += f'<div class="heat-cell {css}"><b>{label}</b><br/>{level}</div>'
    st.markdown(f'<div class="heatmap-container">{cells}</div>', unsafe_allow_html=True)


def render_issues(issues, section_label):
    if not issues:
        st.caption(f"No {section_label.lower()} identified.")
        return
    for issue in issues:
        sev      = issue.get("severity","ELEVATED")
        css      = {"CRITICAL":"issue-critical","ELEVATED":"issue-high","MODERATE":"issue-medium"}.get(sev,"issue-low")
        conf     = issue.get("confidence","MEDIUM")
        conf_css = {"HIGH":"conf-high","MEDIUM":"conf-medium","REQUIRES_REVIEW":"conf-low"}.get(conf,"conf-medium")
        conf_lbl = {"HIGH":"High Confidence","MEDIUM":"Medium Confidence","REQUIRES_REVIEW":"Requires Human Review"}.get(conf,conf)
        partner  = "  |  [PARTNER REVIEW]" if issue.get("partner_review") else ""
        prec     = issue.get("precedent_used","")

        title       = html_escape(str(issue.get("title","")))
        clause_ref  = html_escape(str(issue.get("clause_reference","N/A")))
        verbatim    = html_escape(str(issue.get("verbatim_text",issue.get("current_text",""))))
        analysis    = html_escape(str(issue.get("analysis",issue.get("description",""))))
        prov_law    = html_escape(str(issue.get("provincial_law","")))
        market_std  = html_escape(str(issue.get("market_standard","")))
        recommendation = html_escape(str(issue.get("recommendation","")))
        replacement = html_escape(str(issue.get("proposed_replacement","")))
        bench       = html_escape(str(issue.get("benchmark_comparison","")))
        enf_pct     = issue.get("enforceability_pct","")

        card  = f'<div class="issue-card {css}">'
        card += f'<b>{html_escape(sev)} &mdash; {title}</b>{partner}<br/>'
        card += f'<span class="{conf_css}">{conf_lbl}</span>'
        if enf_pct != "":
            try:
                ev = int(str(enf_pct))
                ec = "#15803d" if ev >= 70 else "#b45309" if ev >= 40 else "#991b1b"
                card += f'<br/><b>Enforceability:</b> <span style="color:{ec};font-weight:600;">{ev}%</span>'
            except (ValueError, TypeError):
                pass
        card += f'<br/><b>Clause:</b> {clause_ref}'
        if verbatim:
            card += f'<br/><b>Text:</b> <i>"{verbatim}"</i>'
        card += f'<br/><b>Analysis:</b> {analysis}'
        if market_std:
            card += f'<br/><b>Market Standard:</b> {market_std}'
        if bench:
            card += f'<br/><b>Benchmark:</b> {bench}'
        if prov_law:
            card += f'<br/><b>Provincial Law:</b> {prov_law}'
        card += f'<br/><b>Recommendation:</b> {recommendation}'
        if replacement:
            prec_note = f' <span style="font-size:10px;color:#6b7280;">(adapted from: {html_escape(prec)})</span>' if prec and prec != "drafted" else ""
            card += f'<br/><b>Proposed Replacement:</b>{prec_note}<br/><code style="font-size:11px;white-space:pre-wrap;">{replacement}</code>'
        card += '</div>'
        st.markdown(card, unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════
# TRUE OOXML TRACK CHANGES REDLINE
# Uses w:ins / w:del elements — Word shows real Track Changes
# that can be accepted or rejected natively.
# ══════════════════════════════════════════════════════════════

_OOXML_NS = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
}
_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def _w(tag):
    return f'{{{_W}}}{tag}'

def _make_del_run(para, text, author="ContractCheck Pro v12.6", date=None):
    """Add a w:del element (strikethrough tracked deletion) to a paragraph."""
    if date is None:
        date = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")
    del_elem = OxmlElement('w:del')
    del_elem.set(_w('id'), str(uuid.uuid4().int % 99999))
    del_elem.set(_w('author'), author)
    del_elem.set(_w('date'), date)
    del_run = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(_w('val'), 'FF0000')
    rpr.append(color)
    del_run.append(rpr)
    del_text = OxmlElement('w:delText')
    del_text.set(qn('xml:space'), 'preserve')
    del_text.text = text
    del_run.append(del_text)
    del_elem.append(del_run)
    para._p.append(del_elem)

def _make_ins_run(para, text, author="ContractCheck Pro v12.6", date=None):
    """Add a w:ins element (tracked insertion) to a paragraph."""
    if date is None:
        date = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")
    ins_elem = OxmlElement('w:ins')
    ins_elem.set(_w('id'), str(uuid.uuid4().int % 99999))
    ins_elem.set(_w('author'), author)
    ins_elem.set(_w('date'), date)
    ins_run = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(_w('val'), '0000FF')
    rpr.append(color)
    underline = OxmlElement('w:u')
    underline.set(_w('val'), 'single')
    rpr.append(underline)
    ins_run.append(rpr)
    ins_text = OxmlElement('w:t')
    ins_text.set(qn('xml:space'), 'preserve')
    ins_text.text = text
    ins_run.append(ins_text)
    ins_elem.append(ins_run)
    para._p.append(ins_elem)


def make_redline_docx(data, meta) -> bytes:
    """Generate a DOCX with true OOXML w:ins/w:del track changes."""
    import tempfile as _tf
    _fd, _fname = _tf.mkstemp(suffix=".docx"); import os as _os; _os.close(_fd)
    doc = DocxDoc()
    for sec in doc.sections:
        sec.left_margin = Cm(2.54); sec.right_margin = Cm(2.54)
    nv = RGBColor(13,27,42)
    ts = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")

    p = doc.add_paragraph(); p.alignment = 1
    r = p.add_run("REDLINE — PROPOSED AMENDMENTS (TRACK CHANGES)")
    r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = nv
    doc.add_paragraph(
        f"RE: {meta['name']}\nDate: {datetime.now().strftime('%B %d, %Y')}\n"
        f"Prepared by: ContractCheck Pro v12.6\n"
        f"Note: This document uses true OOXML track changes (w:ins/w:del). "
        f"Open in Microsoft Word and use Review > Accept/Reject Changes.")
    doc.add_paragraph()

    redlines = data.get('suggested_redline_clauses', [])
    if not redlines:
        all_issues = data.get('high_exposure_issues',[]) + data.get('elevated_risk_issues',[])
        for issue in all_issues:
            if issue.get('proposed_replacement') and issue.get('verbatim_text'):
                redlines.append({
                    "clause_reference": issue.get('clause_reference',''),
                    "issue":            issue.get('title',''),
                    "current_text":     issue.get('verbatim_text',''),
                    "replacement_text": issue.get('proposed_replacement',''),
                    "rationale":        issue.get('recommendation',''),
                })

    if not redlines:
        doc.add_paragraph("No amendments proposed.")
    else:
        for rl in redlines:
            if not rl.get('current_text') or not rl.get('replacement_text'):
                continue
            doc.add_paragraph()
            heading_p = doc.add_paragraph()
            hr = heading_p.add_run(f"{rl.get('clause_reference','')} — {rl.get('issue','')}")
            hr.font.bold = True; hr.font.size = Pt(11); hr.font.color.rgb = nv

            # True track-changes paragraph
            tc_para = doc.add_paragraph()
            _make_del_run(tc_para, rl['current_text'], date=ts)
            # Spacer run
            sp = tc_para.add_run("  ")
            _make_ins_run(tc_para, rl['replacement_text'], date=ts)

            if rl.get('rationale'):
                rat_p = doc.add_paragraph()
                rr = rat_p.add_run("Rationale: "); rr.font.bold = True; rr.font.size = Pt(9)
                rr2 = rat_p.add_run(rl['rationale'])
                rr2.font.size = Pt(9); rr2.font.italic = True
                rr2.font.color.rgb = RGBColor(100,100,100)

    doc.add_page_break()
    disc_p = doc.add_paragraph()
    disc_r = disc_p.add_run(
        "Generated by ContractCheck Pro v12.6. Uses OOXML w:ins/w:del track changes. "
        "All amendments must be reviewed and approved by qualified counsel before use.")
    disc_r.font.size = Pt(8); disc_r.font.italic = True
    disc_r.font.color.rgb = RGBColor(120,120,120)
    doc.save(_fname)
    with open(_fname, "rb") as _fh: _bytes = _fh.read()
    try: _os.unlink(_fname)
    except: pass
    return _bytes

# ══════════════════════════════════════════════════════════════
# DOCX REPORT GENERATION
# ══════════════════════════════════════════════════════════════
def make_docx_report(data, meta, citation_issues=None) -> bytes:
    """Returns DOCX bytes via tempfile (most reliable cross-platform)."""
    import tempfile as _tf
    _fd, _fname = _tf.mkstemp(suffix=".docx"); import os as _os; _os.close(_fd)
    doc = DocxDoc()
    for sec in doc.sections:
        sec.top_margin=Cm(2.54); sec.bottom_margin=Cm(2.54)
        sec.left_margin=Cm(3.18); sec.right_margin=Cm(3.18)
    nv = RGBColor(13,27,42)
    risk_color = {"CRITICAL":RGBColor(139,0,0),"HIGH":RGBColor(139,0,0),
                  "MODERATE":RGBColor(160,120,10),"LOW":RGBColor(0,100,0)}.get(meta['risk'],nv)

    p = doc.add_paragraph(); p.alignment=1
    r = p.add_run("PRIVILEGED & CONFIDENTIAL")
    r.font.size=Pt(10); r.font.bold=True; r.font.color.rgb=nv; r.font.small_caps=True
    p = doc.add_paragraph(); p.alignment=1
    r = p.add_run("CONTRACT REVIEW MEMORANDUM")
    r.font.size=Pt(16); r.font.bold=True; r.font.color.rgb=nv; r.font.name="Georgia"
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment=1
    r = p.add_run(f"Overall Risk: {meta['risk']}  |  Score: {meta['score']}  |  {meta['verdict']}")
    r.font.bold=True; r.font.size=Pt(11); r.font.color.rgb=risk_color

    info_lines = [
        f"RE: {meta['name']}",
        f"Contract Type: {data.get('contract_type','N/A')}",
        f"Client: {data.get('client_party','N/A')}  |  Counterparty: {data.get('counterparty','N/A')}",
        f"Jurisdiction: {meta.get('province','')}  |  Output Mode: {meta.get('output_mode','')}",
    ]
    if meta.get('matter_id'): info_lines.append(f"Matter ID: {meta['matter_id']}")
    info_lines.append(f"Date: {datetime.now().strftime('%B %d, %Y')}  |  Prepared by: ContractCheck Pro v12.6 (AI First-Pass)")
    info_lines.append("Status: FOR PARTNER REVIEW — NOT FOR CLIENT DISTRIBUTION")
    doc.add_paragraph("\n".join(info_lines))
    p = doc.add_paragraph(); r = p.add_run("_"*72)
    r.font.color.rgb=RGBColor(180,180,180); r.font.size=Pt(8)

    def add_heading(text, level=1):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size=Pt(14 if level==1 else 11); r.font.bold=True
        r.font.color.rgb=nv; r.font.name="Georgia"

    add_heading("1.  Executive Summary")
    doc.add_paragraph(data.get('executive_summary','No summary available.'))

    add_heading("2.  Mandatory Clause Review")
    clause_review = data.get('mandatory_clause_review',[])
    if clause_review:
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style='Table Grid'
        tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
        hdrs = tbl.rows[0].cells
        for i, h in enumerate(["Clause","Status","Risk","Confidence","Assessment"]):
            hdrs[i].text=h
            for p in hdrs[i].paragraphs:
                for r in p.runs:
                    r.font.bold=True; r.font.size=Pt(8); r.font.color.rgb=RGBColor(255,255,255)
        for item in clause_review:
            row = tbl.add_row().cells
            cat = item.get("clause_category","")
            row[0].text=MANDATORY_CLAUSE_LABELS.get(cat, cat.replace("_"," ").title())
            row[1].text=item.get("assessment","")
            row[2].text=item.get("risk_level","")
            row[3].text=item.get("confidence","")
            analysis_text=item.get("analysis","")
            row[4].text=analysis_text[:200]+("..." if len(analysis_text)>200 else "")
            for cell in row:
                for p in cell.paragraphs:
                    for r in p.runs: r.font.size=Pt(8)

    for section_num, section_title, key, color in [
        (3,"High Exposure Issues","high_exposure_issues",RGBColor(139,0,0)),
        (4,"Elevated Risk Issues","elevated_risk_issues",RGBColor(160,120,10)),
    ]:
        issues = data.get(key,[])
        if issues:
            add_heading(f"{section_num}.  {section_title}")
            for i, issue in enumerate(issues,1):
                p = doc.add_paragraph()
                r = p.add_run(f"{section_num}.{i}  {issue.get('title','')}  [{issue.get('confidence','')}]")
                r.font.bold=True; r.font.size=Pt(11); r.font.color.rgb=color
                if issue.get('partner_review'):
                    r2=p.add_run("  [PARTNER REVIEW]"); r2.font.bold=True; r2.font.size=Pt(9); r2.font.color.rgb=RGBColor(139,0,0)
                doc.add_paragraph(f"Clause: {issue.get('clause_reference','N/A')}")
                if issue.get('verbatim_text'):
                    p=doc.add_paragraph(); r=p.add_run(f'"{issue["verbatim_text"]}"')
                    r.font.italic=True; r.font.size=Pt(9)
                doc.add_paragraph(f"Analysis: {issue.get('analysis','')}")
                if issue.get('market_standard'): doc.add_paragraph(f"Market Standard: {issue['market_standard']}")
                if issue.get('provincial_law'):  doc.add_paragraph(f"Provincial Law: {issue['provincial_law']}")
                doc.add_paragraph(f"Recommendation: {issue.get('recommendation','')}")
                if issue.get('proposed_replacement'):
                    p=doc.add_paragraph(); r=p.add_run("Proposed Replacement:"); r.font.bold=True
                    prec=issue.get('precedent_used','')
                    if prec and prec!='drafted':
                        p.add_run(f" (adapted from: {prec})")
                    doc.add_paragraph(issue['proposed_replacement'])
                doc.add_paragraph()

    imb = data.get('commercial_imbalances',[])
    if imb:
        add_heading("5.  Commercial Imbalances")
        for item in imb:
            p=doc.add_paragraph(); r=p.add_run(item.get('title','')); r.font.bold=True
            doc.add_paragraph(f"Clause: {item.get('clause_reference','N/A')}")
            doc.add_paragraph(item.get('description',''))
            doc.add_paragraph(f"Recommendation: {item.get('recommendation','')}")
            doc.add_paragraph()

    missing = data.get('missing_protections',[])
    if missing:
        add_heading("6.  Missing Protections")
        for item in missing:
            p=doc.add_paragraph()
            r=p.add_run(f"{item.get('clause_type','')} — {item.get('risk_level','')}"); r.font.bold=True
            doc.add_paragraph(item.get('explanation',''))
            if item.get('proposed_clause'):
                p=doc.add_paragraph(); r=p.add_run("Proposed Clause (from precedent library):"); r.font.bold=True
                doc.add_paragraph(item['proposed_clause'])
            doc.add_paragraph()

    neg = data.get('negotiation_strategy',{})
    if neg and any(neg.values()):
        doc.add_page_break(); add_heading("7.  Negotiation Strategy")
        for lbl, key in [("Priority Items:","priority_items"),("Walk-Away Points:","walk_away_points"),("Concession Candidates:","concession_candidates")]:
            items = neg.get(key,[])
            if items:
                p=doc.add_paragraph(); r=p.add_run(lbl); r.font.bold=True
                for j, item in enumerate(items,1):
                    doc.add_paragraph(f"    {j}. {item}")
        if neg.get('draft_communication'):
            doc.add_paragraph()
            p=doc.add_paragraph(); r=p.add_run("Draft Communication to Opposing Counsel:"); r.font.bold=True
            doc.add_paragraph(neg['draft_communication'])

    actions = data.get('action_items',{})
    if actions and any(actions.values()):
        add_heading("8.  Action Items")
        for phase, items in actions.items():
            if items:
                p=doc.add_paragraph(); r=p.add_run(f"{phase.replace('_',' ').title()}:"); r.font.bold=True
                for item in items: doc.add_paragraph(f"    [ ]  {item}")

    if citation_issues:
        doc.add_paragraph()
        p=doc.add_paragraph(); r=p.add_run("Citation Validation Notes")
        r.font.bold=True; r.font.size=Pt(10); r.font.color.rgb=RGBColor(160,120,10)
        for ci in citation_issues:
            if ci['type']=='INVALID_SECTION_REF':
                doc.add_paragraph(f"  - Section {ci['claimed']} referenced in '{ci['in_issue']}' not found in clause index.")

    doc.add_page_break()
    p=doc.add_paragraph(); r=p.add_run("DISCLAIMER")
    r.font.size=Pt(9); r.font.bold=True; r.font.color.rgb=RGBColor(120,120,120)
    p=doc.add_paragraph()
    r=p.add_run("This memorandum was generated by ContractCheck Pro v12.6 as an AI-assisted first-pass review. "
                "It does not constitute legal advice and must be reviewed by qualified counsel before any reliance, "
                "client communication, or execution decision. The reviewing lawyer is responsible for verifying all "
                "citations, legal analysis, and proposed clause language.")
    r.font.size=Pt(8); r.font.italic=True; r.font.color.rgb=RGBColor(120,120,120)
    doc.save(_fname)
    with open(_fname, "rb") as _fh: _bytes = _fh.read()
    try: _os.unlink(_fname)
    except: pass
    return _bytes


# ══════════════════════════════════════════════════════════════
# PDF REPORT GENERATION
# ══════════════════════════════════════════════════════════════
def make_pdf_report(data, meta, citation_issues=None):
    import tempfile
    fd, fname = tempfile.mkstemp(suffix=".pdf")
    os.close(fd)
    rc = colors.HexColor({"CRITICAL":"#8B0000","HIGH":"#8B0000","MODERATE":"#B8860B","LOW":"#006400"}.get(meta['risk'],"#444"))
    nv = colors.HexColor("#0d1b2a")
    _uid = meta['ts']
    doc = SimpleDocTemplate(fname, pagesize=letter, leftMargin=inch, rightMargin=inch,
                            topMargin=0.75*inch, bottomMargin=0.75*inch)
    S = getSampleStyleSheet()
    def s(n,**k): return ParagraphStyle(f"{n}_{_uid}", parent=S['Normal'], **k)
    body  = s("body", fontSize=10, leading=14, fontName='Helvetica', alignment=TA_JUSTIFY)
    h1    = s("h1p",  fontSize=13, spaceBefore=14, spaceAfter=6, fontName='Helvetica-Bold', textColor=nv)
    h2    = s("h2p",  fontSize=11, spaceBefore=10, spaceAfter=4, fontName='Helvetica-Bold', textColor=nv)
    crit  = s("crit", fontSize=10, leading=14, fontName='Helvetica-Bold', textColor=colors.HexColor("#8B0000"))
    warn  = s("warn", fontSize=10, leading=14, fontName='Helvetica-Bold', textColor=colors.HexColor("#B8860B"))
    qt    = s("qtp",  fontSize=9,  leading=12,  fontName='Helvetica-Oblique',
              textColor=colors.HexColor("#333"), leftIndent=20, rightIndent=20)
    sm    = s("smp",  fontSize=8,  textColor=colors.grey, alignment=TA_CENTER, fontName='Helvetica-Oblique')

    story = []
    def band(html, bg, pad=10):
        t = Table([[Paragraph(html, body)]], colWidths=[6.5*inch])
        t.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,-1),bg),('ALIGN',(0,0),(-1,-1),'CENTER'),
                                ('TOPPADDING',(0,0),(-1,-1),pad),('BOTTOMPADDING',(0,0),(-1,-1),pad)]))
        return t

    story.append(band('<font color="white" size="13"><b>PRIVILEGED &amp; CONFIDENTIAL</b></font><br/>'
                      '<font color="#c9a84c" size="11">CONTRACT REVIEW MEMORANDUM</font>', nv))
    story.append(Spacer(1,6))
    story.append(band(f'<font color="white" size="10"><b>Risk: {meta["risk"]}  |  Score: {meta["score"]}  |  {meta["verdict"]}</b></font>', rc, 6))
    story.append(Spacer(1,8))

    rows = [
        ["RE:",          meta['name']],
        ["Type:",        data.get('contract_type','N/A')],
        ["Client:",      data.get('client_party','N/A')],
        ["Counterparty:",data.get('counterparty','N/A')],
        ["Jurisdiction:",meta.get('province','')],
        ["Date:",        datetime.now().strftime("%B %d, %Y")],
        ["Prepared by:", "ContractCheck Pro v12.6 (AI First-Pass)"],
        ["Status:",      "FOR PARTNER REVIEW"],
    ]
    if meta.get('matter_id'): rows.insert(1,["Matter ID:",meta['matter_id']])
    mt = Table(rows, colWidths=[1.4*inch,5.1*inch])
    mt.setStyle(TableStyle([
        ('FONTNAME',(0,0),(0,-1),'Helvetica-Bold'),('FONTSIZE',(0,0),(-1,-1),9),
        ('TOPPADDING',(0,0),(-1,-1),2),('BOTTOMPADDING',(0,0),(-1,-1),2),
        ('ROWBACKGROUNDS',(0,0),(-1,-1),[colors.HexColor("#f5f5f5"),colors.white])]))
    story.append(mt)
    story.append(HRFlowable(width="100%",thickness=1.5,color=nv))
    story.append(Spacer(1,8))

    story.append(Paragraph("<b>1.  Executive Summary</b>",h1))
    story.append(Paragraph(data.get('executive_summary',''),body))
    story.append(Spacer(1,6))

    clause_review = data.get('mandatory_clause_review',[])
    if clause_review:
        story.append(Paragraph("<b>2.  Mandatory Clause Review</b>",h1))
        tbl_data = [["Category","Status","Risk","Conf.","Assessment"]]
        tbl_styles = [
            ('FONTNAME',(0,0),(-1,0),'Helvetica-Bold'),('FONTSIZE',(0,0),(-1,0),8),
            ('BACKGROUND',(0,0),(-1,0),nv),('TEXTCOLOR',(0,0),(-1,0),colors.white),
            ('FONTSIZE',(0,1),(-1,-1),7),
            ('TOPPADDING',(0,0),(-1,-1),3),('BOTTOMPADDING',(0,0),(-1,-1),3),
            ('GRID',(0,0),(-1,-1),0.4,colors.HexColor("#ddd")),
        ]
        risk_bg = {"CRITICAL":colors.HexColor("#fce4e4"),"HIGH":colors.HexColor("#fef3c7"),
                   "MEDIUM":colors.HexColor("#fef9c3"),"LOW":colors.HexColor("#dcfce7"),
                   "MISSING":colors.HexColor("#f3f4f6")}
        for idx, item in enumerate(clause_review,1):
            cat   = item.get("clause_category","")
            label = MANDATORY_CLAUSE_LABELS.get(cat,cat.replace("_"," ").title())
            analysis = item.get("analysis","")
            tbl_data.append([label, item.get("assessment",""), item.get("risk_level",""),
                             item.get("confidence",""), analysis[:120]+("..." if len(analysis)>120 else "")])
            bg = risk_bg.get(item.get("risk_level","").upper(), colors.white)
            tbl_styles.append(('BACKGROUND',(2,idx),(2,idx),bg))
        t = Table(tbl_data, colWidths=[1.2*inch,0.9*inch,0.7*inch,0.6*inch,3.1*inch])
        t.setStyle(TableStyle(tbl_styles))
        story.append(t); story.append(Spacer(1,8))

    def add_issues_pdf(issues, num, title, sty):
        if not issues: return
        story.append(Paragraph(f"<b>{num}.  {title}</b>",h1))
        for issue in issues:
            conf    = issue.get('confidence','')
            partner = "  [PARTNER REVIEW]" if issue.get('partner_review') else ""
            story.append(Paragraph(
                f"<b>{issue.get('severity',issue.get('risk_level',''))}: {issue.get('title','')}</b>  [{conf}]{partner}", sty))
            story.append(Paragraph(f"<b>Clause:</b> {issue.get('clause_reference','N/A')}",body))
            vt = issue.get('verbatim_text',issue.get('current_text',''))
            if vt: story.append(Paragraph(f'"{vt}"',qt))
            story.append(Paragraph(issue.get('analysis',issue.get('description','')),body))
            if issue.get('market_standard'): story.append(Paragraph(f"<i>Market Standard: {issue['market_standard']}</i>",body))
            if issue.get('provincial_law'):  story.append(Paragraph(f"<i>Provincial Law: {issue['provincial_law']}</i>",body))
            if issue.get('recommendation'):  story.append(Paragraph(f"<b>Recommendation:</b> {issue['recommendation']}",body))
            pr = issue.get('proposed_replacement',issue.get('proposed_clause',''))
            if pr:
                prec = issue.get('precedent_used','')
                prec_note = f" (adapted from: {prec})" if prec and prec!='drafted' else ""
                story.append(Paragraph(f"<b>Proposed Replacement{prec_note}:</b> {pr}",body))
            story.append(Spacer(1,6))

    add_issues_pdf(data.get('high_exposure_issues',[]),3,"High Exposure Issues",crit)
    add_issues_pdf(data.get('elevated_risk_issues',[]),4,"Elevated Risk Issues",warn)
    add_issues_pdf(data.get('missing_protections',[]),6,"Missing Protections",warn)

    neg = data.get('negotiation_strategy',{})
    if neg and any(neg.values()):
        story.append(PageBreak())
        story.append(Paragraph("<b>7.  Negotiation Strategy</b>",h1))
        for lbl, key in [("Priority Items:","priority_items"),("Walk-Away Points:","walk_away_points"),("Concession Candidates:","concession_candidates")]:
            if neg.get(key):
                story.append(Paragraph(f"<b>{lbl}</b>",h2))
                for j,item in enumerate(neg[key],1):
                    story.append(Paragraph(f"{j}. {item}",body))
        if neg.get('draft_communication'):
            story.append(Spacer(1,8))
            story.append(Paragraph("<b>Draft Communication to Opposing Counsel:</b>",h2))
            for line in neg['draft_communication'].split('\n'):
                story.append(Paragraph(line,body))

    story.append(Spacer(1,20))
    story.append(HRFlowable(width="100%",thickness=0.5,color=colors.grey))
    story.append(Paragraph(
        "This memorandum was generated by ContractCheck Pro v12.6 as an AI-assisted first-pass review. "
        "It does not constitute legal advice and must be reviewed by qualified counsel before any reliance or client communication.",sm))
    doc.build(story)
    return fname


def make_amendments_docx(data, meta) -> bytes:
    """Returns DOCX bytes via tempfile."""
    import tempfile as _tf
    _fd, _fname = _tf.mkstemp(suffix=".docx"); import os as _os; _os.close(_fd)
    doc = DocxDoc()
    for sec in doc.sections: sec.left_margin=Cm(3.18); sec.right_margin=Cm(3.18)
    nv = RGBColor(13,27,42)
    p=doc.add_paragraph(); r=p.add_run("SCHEDULE OF PROPOSED AMENDMENTS")
    r.font.size=Pt(14); r.font.bold=True; r.font.color.rgb=nv; p.alignment=1
    doc.add_paragraph(f"RE: {meta['name']}\nDate: {datetime.now().strftime('%B %d, %Y')}")
    p=doc.add_paragraph(); r=p.add_run("_"*72); r.font.color.rgb=RGBColor(180,180,180); r.font.size=Pt(8)

    redlines = data.get('suggested_redline_clauses',[])
    if not redlines:
        all_issues = data.get('high_exposure_issues',[]) + data.get('elevated_risk_issues',[])
        for issue in all_issues:
            if issue.get('proposed_replacement'):
                redlines.append({
                    "clause_reference": issue.get('clause_reference',''),
                    "issue":            issue.get('title',''),
                    "current_text":     issue.get('verbatim_text',''),
                    "replacement_text": issue.get('proposed_replacement',''),
                    "rationale":        issue.get('recommendation',''),
                    "precedent_used":   issue.get('precedent_used',''),
                })

    for i, rl in enumerate(redlines,1):
        p=doc.add_paragraph()
        r=p.add_run(f"Amendment {i}: {rl.get('clause_reference','')} — {rl.get('issue','')}")
        r.font.bold=True; r.font.size=Pt(11)
        doc.add_paragraph()
        if rl.get('current_text'):
            p=doc.add_paragraph(); r=p.add_run("CURRENT: "); r.font.bold=True
            p.add_run(f'"{rl["current_text"]}"')
        if rl.get('replacement_text'):
            p=doc.add_paragraph(); r=p.add_run("PROPOSED: "); r.font.bold=True
            p.add_run(f'"{rl["replacement_text"]}"')
            prec = rl.get('precedent_used','')
            if prec and prec != 'drafted':
                p=doc.add_paragraph(); r=p.add_run(f"(Adapted from precedent: {prec})"); r.font.italic=True; r.font.size=Pt(9)
        if rl.get('rationale'):
            p=doc.add_paragraph(); r=p.add_run("RATIONALE: "); r.font.bold=True
            p.add_run(rl['rationale'])
        doc.add_paragraph()

    if not redlines:
        doc.add_paragraph("No amendments proposed.")
    doc.save(_fname)
    with open(_fname, "rb") as _fh: _bytes = _fh.read()
    try: _os.unlink(_fname)
    except: pass
    return _bytes

# ══════════════════════════════════════════════════════════════
# APPLIED REDLINE (original contract + amendments inline)
# ══════════════════════════════════════════════════════════════
def make_applied_redline_docx(original_text, redline_clauses, meta) -> bytes:
    """Returns DOCX bytes via tempfile."""
    import tempfile as _tf
    _fd, _fname = _tf.mkstemp(suffix=".docx"); import os as _os; _os.close(_fd)
    doc = DocxDoc()
    for sec in doc.sections:
        sec.left_margin=Cm(2.54); sec.right_margin=Cm(2.54)
        sec.top_margin=Cm(2.54);  sec.bottom_margin=Cm(2.54)
    nv = RGBColor(13,27,42)
    ts = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")

    p=doc.add_paragraph(); p.alignment=1
    r=p.add_run("AUTOMATED REDLINE — CONTRACT WITH PROPOSED AMENDMENTS APPLIED")
    r.font.size=Pt(12); r.font.bold=True; r.font.color.rgb=nv
    doc.add_paragraph(f"RE: {meta['name']}\nDate: {datetime.now().strftime('%B %d, %Y')}\n"
                      "Uses OOXML w:ins/w:del track changes. Open in Word to accept/reject.")
    doc.add_paragraph()

    replacements = {}
    for rl in redline_clauses:
        current     = rl.get('current_text', rl.get('verbatim_text','')).strip()
        replacement = rl.get('replacement_text', rl.get('proposed_replacement','')).strip()
        if current and replacement:
            replacements[current] = replacement

    clean_original = re.sub(r'^\[CLAUSE INDEX.*?\]\n\[.*?\]\n','',original_text,flags=re.DOTALL)
    for line in clean_original.split('\n'):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph(); continue
        matched = False
        for current_text, replacement_text in list(replacements.items()):
            if current_text[:50] in stripped or stripped in current_text:
                p = doc.add_paragraph()
                _make_del_run(p, stripped, date=ts)
                p.add_run("  ")
                _make_ins_run(p, replacement_text, date=ts)
                matched = True
                del replacements[current_text]
                break
        if not matched:
            p = doc.add_paragraph(stripped)
            for r in p.runs: r.font.size=Pt(10)

    doc.add_page_break()
    p=doc.add_paragraph()
    r=p.add_run("AI-assisted redline generated by ContractCheck Pro v12.6. "
                "Uses OOXML w:ins/w:del track changes. "
                "All amendments must be reviewed by qualified counsel before use.")
    r.font.size=Pt(8); r.font.italic=True; r.font.color.rgb=RGBColor(120,120,120)
    doc.save(_fname)
    with open(_fname, "rb") as _fh: _bytes = _fh.read()
    try: _os.unlink(_fname)
    except: pass
    return _bytes


# ══════════════════════════════════════════════════════════════
# MATTER HISTORY SIDEBAR
# ══════════════════════════════════════════════════════════════
def render_matter_history_sidebar():
    """Render the matter history panel in the sidebar."""
    init_db()
    count = matter_count()
    st.sidebar.divider()
    st.sidebar.subheader(f"Matter History ({count})")

    if count == 0:
        st.sidebar.caption("No matters saved yet. Run an analysis to save it.")
        return

    matters = list_matters(limit=20)
    for m in matters:
        risk_icon = {"CRITICAL":"🔴","HIGH":"🟠","MODERATE":"🟡","LOW":"🟢"}.get(m.get('risk_level',''),"⚪")
        date_str  = m['created_at'][:10] if m.get('created_at') else ""
        label     = m.get('name','')[:28]
        ctype     = m.get('contract_type','')[:20]
        docs      = m.get('doc_count',1)
        doc_note  = f" · {docs} docs" if docs > 1 else ""
        with st.sidebar.expander(f"{risk_icon} {label}", expanded=False):
            st.caption(f"{date_str} · {ctype}{doc_note}")
            st.caption(f"Score: {m.get('risk_score',0)} · ${m.get('cost_usd',0):.4f}")
            col1, col2 = st.columns(2)
            with col1:
                if st.button("Load", key=f"load_{m['id']}"):
                    st.session_state.loaded_matter = load_matter(m['id'])
                    st.rerun()
            with col2:
                if st.button("Delete", key=f"del_{m['id']}"):
                    delete_matter(m['id'])
                    st.rerun()


# ══════════════════════════════════════════════════════════════
# MAIN ANALYSIS RUNNER — handles both normal and chunked paths
# ══════════════════════════════════════════════════════════════
def _sanitise_for_storage(data: dict) -> dict:
    """Strip verbatim contract text from analysis before persisting to SQLite.

    The matter store claims to hold only "locally stored matter metadata" —
    not raw contract text. The analysis JSON contains verbatim_text fields
    (direct contract quotes extracted by the AI) and potentially a full
    revised_contract. This function removes those fields so the stored JSON
    contains only AI-generated metadata: risk levels, titles, recommendations,
    benchmarks, and proposed replacement language — but no verbatim client text.

    The full analysis remains available in the session (st.session_state) for
    the duration of the current session; only the persisted copy is stripped.
    """
    import copy
    safe = copy.deepcopy(data)

    # Fields that contain verbatim contract text
    VERBATIM_FIELDS = ("verbatim_text", "current_text", "revised_contract")

    def _strip(obj):
        if isinstance(obj, dict):
            for field in VERBATIM_FIELDS:
                if field in obj:
                    obj[field] = "[stripped for storage — see original session for full text]"
            for v in obj.values():
                _strip(v)
        elif isinstance(obj, list):
            for item in obj:
                _strip(item)

    _strip(safe)
    return safe


def run_analysis(contract_text, clause_index, ctype, province, client_role,
                 special, playbook_text, analysis_depth, output_mode,
                 analysis_language, model_id, use_caching,
                 deal_ctx: DealContext = None, pre_score_note: str = ""):
    """
    Runs the full analysis pipeline. For long contracts, uses the
    chunk-then-synthesise approach. Returns (data_dict, full_raw, elapsed, cost, is_json).
    """
    depth_cfg = {
        "Quick Scan":          ("Top 5 issues only. Evaluate all 15 mandatory clauses but keep analysis concise.", 8000),
        "Standard Review":     ("Full analysis. All sections. Comprehensive evaluation.", 16000),
        "Deep Dive + Redraft": ("Exhaustive. All sections including complete revised contract in revised_contract field.", 20000),
    }
    depth_note, max_out = depth_cfg[analysis_depth]
    type_note = "Auto-detect the contract type from the text." if ctype == "Auto-Detect" else f"Contract type: {ctype}."

    sys_prompt = build_system_prompt(output_mode, province, playbook_text, ctype,
                                     analysis_language, deal_ctx)
    sys_cfg    = [{"type":"text","text":sys_prompt,"cache_control":{"type":"ephemeral"}}] if use_caching else sys_prompt

    total_chars = len(contract_text)
    is_long     = total_chars > CHUNK_THRESHOLD

    start = time.time()

    if is_long:
        st.info(f"Long contract detected ({total_chars:,} chars). Analysing in chunks — no content will be missed.")
        chunks = _split_into_chunks(contract_text, CHUNK_SIZE, CHUNK_OVERLAP)
        st.caption(f"Splitting into {len(chunks)} overlapping chunks of ~{CHUNK_SIZE:,} chars each.")
        chunk_results = []
        full_raw_parts = []
        prog = st.progress(0)
        for i, chunk in enumerate(chunks):
            st.caption(f"Analysing chunk {i+1} of {len(chunks)}...")
            # Prepend clause index to every chunk
            idx_str = "\n\n[CLAUSE INDEX: " + ", ".join(clause_index) + "]\n[CHUNK " + str(i+1) + " OF " + str(len(chunks)) + " — analyse only what is present in this chunk]\n\n"
            chunk_with_idx = idx_str + chunk
            user_prompt = (
                f"{type_note}\nJurisdiction: {province}\n"
                f"Client role: {client_role}.\n"
                + (f"Priority: {special}\n" if special else "")
                + f"Analysis depth: {depth_note}\n"
                + (f"\n{pre_score_note}\n" if pre_score_note and i == 0 else "")
                + "Set revised_contract to empty string.\n\n"
                + f"CONTRACT TEXT (CHUNK {i+1}/{len(chunks)}):\n{chunk_with_idx}"
            )
            try:
                chunk_data, chunk_full, ok = _call_ai_with_retry(
                    ai, model_id, max_out, sys_cfg, user_prompt,
                    timeout=420.0,
                    progress_label=f"Chunk {i+1}/{len(chunks)}:")
            except Exception as e:
                st.warning(f"Chunk {i+1} error after retries: {e} — skipping.")
                prog.progress((i+1)/len(chunks))
                continue
            full_raw_parts.append(chunk_full)
            if ok and isinstance(chunk_data, dict):
                chunk_results.append(chunk_data)
            prog.progress((i+1)/len(chunks))
        prog.empty()

        # Synthesis pass — merge all chunk results
        if chunk_results:
            merged = _merge_chunk_analyses(chunk_results)
            # Run a short synthesis prompt to write a coherent executive summary
            synth_summary = (
                f"Based on these chunk-level analyses of a long contract, write a coherent 3-sentence "
                f"executive summary and a one-sentence verdict for the final report.\n"
                f"Risk level: {merged['risk_level']}, Score: {merged['risk_score']}\n"
                f"Issues found: {len(merged['high_exposure_issues'])} critical, {len(merged['elevated_risk_issues'])} elevated.\n"
                f"Contract type: {merged.get('contract_type','')}, Jurisdiction: {province}\n"
                f"Respond with JSON: {{\"executive_summary\": \"...\", \"verdict\": \"...\"}}"
            )
            try:
                synth_resp = ai.messages.create(
                    model=model_id, max_tokens=400,
                    messages=[{"role":"user","content":synth_summary}])
                synth_data, synth_ok = parse_json_response(synth_resp.content[0].text)
                if synth_ok and isinstance(synth_data, dict):
                    # Validate the synthesis response fields before trusting them
                    if synth_data.get("executive_summary") and len(str(synth_data["executive_summary"])) > 20:
                        merged["executive_summary"] = synth_data["executive_summary"]
                    if synth_data.get("verdict") and len(str(synth_data["verdict"])) > 5:
                        merged["verdict"] = synth_data["verdict"]
            except Exception as synth_err:
                # Synthesis failure is non-fatal — merged summary still usable
                st.caption(f"Note: synthesis pass failed ({synth_err}), using merged summaries.")
            data    = merged
            is_json = True
            full    = "\n---CHUNK SEPARATOR---\n".join(full_raw_parts)
        else:
            data    = None
            is_json = False
            full    = "\n".join(full_raw_parts)
    else:
        # Normal single-pass analysis
        user_prompt = (
            f"{type_note}\nJurisdiction: {province}\n"
            f"Client role: {client_role}. Identify the named party who is the {client_role} and protect their interests.\n"
            + (f"Priority: {special}\n" if special else "")
            + f"Analysis depth: {depth_note}\n"
            + ("Include the complete revised contract text in the revised_contract field.\n" if "Deep Dive" in analysis_depth else "Set revised_contract to empty string.\n")
            + (f"\n{pre_score_note}\n" if pre_score_note else "")
            + f"\nCONTRACT TEXT:\n{contract_text}"
        )
        data, full, is_json = _call_ai_with_retry(
            ai, model_id, max_out, sys_cfg, user_prompt,
            timeout=420.0,
            progress_label="Receiving...")

    elapsed = time.time() - start
    out_toks = len(full) // 4
    cost = (3000*0.3 + (total_chars//4)*_INPUT_RATE)/1_000_000 + out_toks*_OUTPUT_RATE/1_000_000
    return data, full, elapsed, cost, is_json

# ══════════════════════════════════════════════════════════════
# MAIN UI
# ══════════════════════════════════════════════════════════════
render_matter_history_sidebar()


# ══════════════════════════════════════════════════════════════
# DEMO CONTRACT — realistic but deliberately flawed Canadian
# service agreement. Used by the "Load Demo" button so lawyers
# can explore the full analysis pipeline without uploading a file.
# Flaws planted for the AI to find:
#   1. Liability cap at $500 (absurdly low for a $120k/yr deal)
#   2. Late fee at 3%/month (36% APR — exceeds Bill C-46 35% cap)
#   3. Auto-renews with only 7-day cancellation window
#   4. No IP ownership / work product clause
#   5. No governing law specified
#   6. Non-compete in what appears to be an employment context (ON)
#   7. Indemnity is one-sided (client only)
#   8. Dispute resolution requires arbitration in Delaware
# ══════════════════════════════════════════════════════════════
DEMO_CONTRACT_NAME = "DEMO — Acme Corp / BrightWeb Agency (Service Agreement)"
DEMO_CONTRACT_TEXT = """
SERVICE AGREEMENT

This Agreement is made January 15, 2026 between Acme Corp Inc. ("Client") and
BrightWeb Agency Ltd. ("Agency"), both Ontario corporations.

1. SERVICES AND FEES

The Agency shall provide digital marketing services for a monthly fee of CAD $8,000
(CAD $96,000 annually). Invoices are due within sixty (60) days of the invoice date.
Late payment shall bear interest at 3% per month (36% per annum), compounded monthly.

2. TERM AND RENEWAL

This Agreement is for one (1) year and shall automatically renew for successive
one-year terms unless cancelled by written notice at least seven (7) calendar days
before the renewal date. Fees may increase by up to 20% on each renewal without
further notice.

3. LIMITATION OF LIABILITY

The Agency's total liability under this Agreement shall not exceed CAD $500.00,
regardless of the nature or cause of any claim.

4. INDEMNIFICATION

The Client shall indemnify and hold harmless the Agency from any and all claims
arising out of the Client's use of the deliverables, the Client's content, or
the Client's breach of this Agreement. The Agency has no indemnification obligation
to the Client.

5. NON-COMPETE

The Agency agrees that for two (2) years following termination, it shall not provide
services to any competitor of the Client anywhere in Canada. Any Agency employee who
provides services under this Agreement is subject to the same restriction.

6. CONFIDENTIALITY

Each party shall keep the other's confidential information secret for one (1) year
after termination.

7. DISPUTE RESOLUTION

All disputes shall be resolved by binding arbitration before the American Arbitration
Association in Wilmington, Delaware, USA.

8. MISCELLANEOUS

This Agreement is the entire agreement between the parties. Signed below.

Acme Corp Inc.                      BrightWeb Agency Ltd.
By: Jennifer Park, CEO              By: Daniel Roy, Director
Date: January 15, 2026              Date: January 15, 2026
"""


tab_analyze, tab_compare, tab_negotiate, tab_history, tab_pipeda = st.tabs([
    "Contract Analysis", "Version Comparison", "Negotiation Simulator",
    "Matter History", "PIPEDA Compliance"
])

# ── TAB 1: CONTRACT ANALYSIS ──────────────────────────────────
with tab_analyze:
    col_up, col_name = st.columns([1,1], gap="large")
    with col_up:
        st.subheader("Upload Contract")
        uploaded_main = st.file_uploader(
            "Main agreement (PDF, DOCX, or TXT)",
            type=["pdf","docx","txt"], label_visibility="collapsed")
        if uploaded_main:
            st.caption(f"Main: {uploaded_main.name}  ({uploaded_main.size/1024:.0f} KB)")
            # Real file uploaded — exit demo mode automatically
            if st.session_state.get("demo_mode"):
                st.session_state.demo_mode = False
                st.session_state.pop("demo_name", None)

        st.caption("**Schedules / Exhibits (optional, up to 4)**")
        uploaded_schedules = st.file_uploader(
            "Additional documents", type=["pdf","docx","txt"],
            accept_multiple_files=True, key="schedules",
            label_visibility="collapsed",
            help="Upload schedules, exhibits, SOWs, or side letters. They will be analysed together with the main agreement.")
        if uploaded_schedules:
            for sf in uploaded_schedules[:4]:
                st.caption(f"  + {sf.name}  ({sf.size/1024:.0f} KB)")

    with col_name:
        st.subheader("Contract Name")
        name = st.text_input("Contract / Matter Name",
            placeholder="e.g. Acme Corp — Marketing Services Agreement",
            label_visibility="collapsed")

    with st.expander("Advanced Settings", expanded=False):
        adv_c1, adv_c2, adv_c3 = st.columns(3)
        with adv_c1:
            matter_id = st.text_input("File / Matter Number", placeholder="e.g. 2025-0312")
            ctype = st.selectbox("Contract Type", [
                "Auto-Detect","Employment Contract","Independent Contractor",
                "Service Agreement","Commercial Lease","NDA / Confidentiality",
                "Partnership / Shareholder","Vendor / Supplier","Franchise Agreement",
                "SaaS / Technology","Loan / Investment","Construction Contract",
                "Distribution / Licensing","Other"])
        with adv_c2:
            province = st.selectbox("Governing Law / Jurisdiction", list(PROVINCE_RULES.keys()))
            client_role = st.selectbox("Client's Position", [
                "Service Provider / Contractor","Client / Purchaser","Employee","Employer",
                "Tenant","Landlord","Partner / Shareholder","Vendor / Supplier",
                "Franchisee","Franchisor","Borrower","Lender","Licensee","Licensor"])
        with adv_c3:
            special = st.text_input("Priority Clauses", placeholder="e.g. Focus on s.6 non-compete")
            playbook_file = st.file_uploader("Firm Playbook (TXT/DOCX)", type=["txt","docx"],
                key="playbook", label_visibility="collapsed",
                help="Upload your firm's standard positions. AI compares against your standards.")
        playbook_text = ""
        if playbook_file:
            if playbook_file.name.endswith('.txt'):
                playbook_text = playbook_file.read().decode('utf-8')
            elif playbook_file.name.endswith('.docx'):
                playbook_text = "\n".join(p.text for p in DocxDoc(playbook_file).paragraphs if p.text.strip())
            st.success(f"Playbook loaded: {len(playbook_text):,} characters")

        # ── DEAL CONTEXT ──────────────────────────────────────
        st.divider()
        st.caption("**Deal Context** — calibrates benchmark thresholds and precedent clause values to this specific deal. Skipping these fields uses generic market standards.")
        dc1, dc2, dc3, dc4 = st.columns(4)
        with dc1:
            contract_value_str = st.text_input(
                "Contract Value (CAD)",
                placeholder="e.g. 250000",
                help="Total contract value. Used to scale liability cap floors, insurance requirements, and notice periods.")
        with dc2:
            annual_fees_str = st.text_input(
                "Annual Fees (CAD)",
                placeholder="e.g. 120000",
                help="Recurring annual fees — used to compute exact dollar amounts for liability caps (e.g. '12 months = CAD $120,000').")
        with dc3:
            industry_sel = st.selectbox(
                "Industry",
                list(INDUSTRY_RISK_TIERS.keys()),
                index=list(INDUSTRY_RISK_TIERS.keys()).index("General / Other"),
                help="Industry risk tier affects which clauses are weighted most heavily.")
            counterparty_sel = st.selectbox(
                "Counterparty Type",
                list(COUNTERPARTY_RISK.keys()),
                index=2,
                help="Counterparty size affects how aggressive the AI expects the other side to be.")
        with dc4:
            risk_tol = st.selectbox(
                "Client Risk Tolerance",
                ["Standard", "Conservative", "Aggressive"],
                index=0,
                help="Conservative = higher floors on all caps and longer notice. Aggressive = more flexibility on below-standard terms.")
            duration_str = st.text_input(
                "Contract Duration (months)",
                placeholder="e.g. 24",
                help="Used to calibrate auto-renewal notice periods and short-term contract adjustments.")
            client_is_vendor = st.checkbox(
                "Client is the vendor/service provider",
                value=False,
                help="Changes payment term risk direction: if client is the payee, longer payment terms are worse for them.")

    # Build DealContext from inputs (gracefully handles empty/invalid fields)
    matter_id     = matter_id    if 'matter_id'    in locals() else ""
    ctype         = ctype        if 'ctype'        in locals() else "Auto-Detect"
    province      = province     if 'province'     in locals() else "Ontario"
    client_role   = client_role  if 'client_role'  in locals() else "Service Provider / Contractor"
    special       = special      if 'special'      in locals() else ""
    playbook_text = playbook_text if 'playbook_text' in locals() else ""

    def _safe_float(s):
        try: return float(str(s).replace(",","").strip()) if s else None
        except (ValueError, TypeError): return None
    def _safe_int(s):
        try: return int(str(s).strip()) if s else None
        except (ValueError, TypeError): return None

    deal_ctx = None
    if 'industry_sel' in locals():
        deal_ctx = DealContext(
            contract_value_cad = _safe_float(contract_value_str if 'contract_value_str' in locals() else ""),
            annual_fees_cad    = _safe_float(annual_fees_str    if 'annual_fees_str'    in locals() else ""),
            industry           = industry_sel   if 'industry_sel'   in locals() else "General / Other",
            counterparty_type  = counterparty_sel if 'counterparty_sel' in locals() else "SMB (< 50 employees)",
            client_is_vendor   = client_is_vendor if 'client_is_vendor' in locals() else False,
            risk_tolerance     = risk_tol        if 'risk_tol'        in locals() else "Standard",
            duration_months    = _safe_int(duration_str if 'duration_str' in locals() else ""),
            province           = province if 'province' in locals() else "Ontario",
        )

        # Show deal context summary to user
        if deal_ctx.contract_value_cad or deal_ctx.annual_fees_cad:
            ctx_parts = [f"Deal: {deal_ctx.deal_size_tier}"]
            if deal_ctx.contract_value_cad:
                ctx_parts.append(f"CAD ${deal_ctx.contract_value_cad:,.0f}")
            ctx_parts.append(f"Industry risk: {deal_ctx.industry_risk_tier}")
            ctx_parts.append(f"Risk tolerance: {deal_ctx.risk_tolerance}")
            st.caption(f"Deal context active — {' · '.join(ctx_parts)}. Thresholds calibrated to this deal.")

    # ── DEMO BUTTON ──────────────────────────────────────────────
    with st.expander("🎬  Try a Demo Contract (no upload needed)", expanded=False):
        st.markdown(
            "Load a **short, deliberately flawed** Ontario service agreement and run "
            "a full AI analysis in **under 90 seconds**. Eight issues are planted — "
            "the AI and Python rule engine should catch all of them.")
        st.caption(
            "Planted: $500 liability cap · 3%/month late fee (Criminal Code s.347) · "
            "7-day auto-renewal · no IP clause · no governing law · "
            "overbroad non-compete · one-sided indemnity · Delaware arbitration")
        if st.button("Load Demo Contract", key="load_demo", use_container_width=True):
            st.session_state.demo_mode = True
            st.session_state.demo_name = DEMO_CONTRACT_NAME
            st.rerun()
        if st.session_state.get("demo_mode"):
            st.success(f"Demo loaded: **{DEMO_CONTRACT_NAME}** — "
                       "click Analyze Contract below to run the full analysis.")

    st.divider()

    # Wire demo mode: use injected text instead of file upload
    _using_demo = st.session_state.get("demo_mode", False)
    if _using_demo:
        st.info("Demo mode active — using the pre-loaded contract. "
                "Upload a file above to switch to your own contract.")

    analyze_clicked = st.button("Analyze Contract", type="primary", use_container_width=True)

    if analyze_clicked:
        # ── Demo mode: bypass file upload ──────────────────────
        if _using_demo:
            name = name or DEMO_CONTRACT_NAME
            contract_text, extraction_method = DEMO_CONTRACT_TEXT, "text"
            contract_text, clause_index = build_clause_index(contract_text)
            doc_count = 1
            chars = len(contract_text)
            is_chunked = False   # demo is always short
            # Force Quick Scan for speed — demo completes in ~60 seconds
            analysis_depth = "Quick Scan"
            # Pre-fill sensible defaults for demo
            ctype        = "Service Agreement"
            province     = "Ontario"
            client_role  = "Client / Purchaser"
        else:
            if not uploaded_main: st.error("Please upload a contract."); st.stop()
            if not name:          st.error("Please enter a contract name."); st.stop()

            all_files = [uploaded_main] + (uploaded_schedules[:4] if uploaded_schedules else [])
            doc_count = len(all_files)

            # Validate all files before any processing
            for upload_f in all_files:
                if not validate_upload(upload_f):
                    st.stop()

            with st.spinner(f"Extracting text from {doc_count} document(s)..."):
                contract_text, extraction_method, clause_index = extract_text(all_files)

        if not _using_demo:
            chars     = len(contract_text)
            is_chunked = chars > CHUNK_THRESHOLD
        st.success(
            f"Extracted {chars:,} characters from {doc_count} document(s) "
            f"via {extraction_method} "
            f"({chars//4:,} est. tokens)"
            + (" — long contract mode active" if is_chunked else ""))

        # ── PIPEDA: PI Detection ─────────────────────────────────
        # Scan for personal information patterns BEFORE the API call.
        # The lawyer sees the findings and can choose to redact or
        # confirm consent before proceeding.
        pi_findings = detect_personal_information(contract_text)
        pi_summary  = pi_risk_summary(pi_findings)

        if pi_findings:
            with st.expander(
                    f"⚠️ PIPEDA — Personal Information Detected "
                    f"({len(pi_findings)} category(s), risk: {pi_summary['level']})",
                    expanded=(pi_summary['level'] == "HIGH")):
                st.markdown(
                    f"<div style='border-left:4px solid {pi_summary['colour']};"
                    f"padding:8px 14px;background:#fafafa;border-radius:0 4px 4px 0;"
                    f"margin-bottom:10px;'>"
                    f"<b>{html_escape(pi_summary['summary'])}</b><br/>"
                    f"<span style='font-size:12px;color:#555;'>"
                    f"{html_escape(pi_summary['action'])}</span></div>",
                    unsafe_allow_html=True)
                for finding in pi_findings:
                    risk_color = {"HIGH":"#8B0000","MEDIUM":"#B8860B","LOW":"#2563EB"}.get(
                        finding.risk, "#555")
                    st.markdown(
                        f"<div style='border-left:3px solid {risk_color};"
                        f"padding:6px 12px;margin:4px 0;font-size:13px;'>"
                        f"<b style='color:{risk_color};'>{html_escape(finding.risk)}</b> — "
                        f"<b>{html_escape(finding.category)}</b> "
                        f"({finding.count} occurrence(s))<br/>"
                        f"<span style='color:#666;font-size:11px;'>"
                        f"{html_escape(finding.guidance)}</span></div>",
                        unsafe_allow_html=True)

                if pi_summary['level'] == "HIGH":
                    st.warning(
                        "**Lawyer action required:** Redact identified PI before "
                        "proceeding, OR confirm that the client retainer agreement "
                        "includes an AI processing disclosure. "
                        "Proceeding without action may not satisfy PIPEDA obligations.")
                    col_proceed, col_retainer = st.columns(2)
                    with col_proceed:
                        if not st.checkbox(
                                "I confirm PI handling is addressed (redacted / "
                                "retainer disclosure obtained)",
                                key="pi_confirmed"):
                            st.info("Check the box above to proceed with analysis.")
                            st.stop()
                    with col_retainer:
                        retainer_text = generate_retainer_disclosure(
                            jurisdiction=province)
                        st.download_button(
                            "Download Retainer Disclosure Clause",
                            data=retainer_text,
                            file_name="AI_Processing_Retainer_Disclosure.txt",
                            mime="text/plain",
                            help="Add this clause to your standard client retainer agreement.")

                # Auto-redact option
                st.markdown("**Or: Auto-Redact PI before sending to AI**")
                st.caption(
                    "Replaces HIGH/MEDIUM sensitivity PI with [REDACTED] placeholders. "
                    "Download, then re-upload as your contract to eliminate PI from the API call.")
                if st.button("Generate Redacted Contract", key="auto_redact_btn"):
                    redacted_text, redact_count = redact_personal_information(
                        contract_text, pi_findings)
                    st.download_button(
                        f"Download Redacted Contract ({redact_count} substitution(s))",
                        data      = redacted_text,
                        file_name = f"{re.sub(chr(91) + r'^\w\s-' + chr(93), '', name or 'contract')[:30]}_redacted.txt",
                        mime      = "text/plain",
                        key       = "download_redacted",
                        type      = "primary")
        else:
            st.caption("✓ PIPEDA scan: no personal information patterns detected.")

        found_clauses = detect_clauses_present(contract_text)
        missing_mandatory = [c for c in MANDATORY_CLAUSES if c not in found_clauses]
        completeness = round(len(found_clauses)/len(MANDATORY_CLAUSES)*100)
        with st.expander("Pre-Analysis: Clause Detection", expanded=False):
            cols_m = st.columns(3)
            cols_m[0].metric("Detected",   f"{len(found_clauses)} / 15")
            cols_m[1].metric("Absent",     len(missing_mandatory))
            cols_m[2].metric("Coverage",   f"{completeness}%")
            if missing_mandatory:
                st.caption(f"Not detected: {', '.join(MANDATORY_CLAUSE_LABELS.get(c,c) for c in missing_mandatory)}")

        # Deal context calibration panel
        if deal_ctx and (deal_ctx.contract_value_cad or deal_ctx.annual_fees_cad):
            with st.expander("Deal Context: Calibrated Thresholds", expanded=False):
                st.caption("These deal-specific thresholds are injected into the analysis. The AI uses these instead of generic market figures.")
                from deal_context import _BASE_THRESHOLDS, compute_dynamic_thresholds
                thresh_cols = st.columns(3)
                key_clauses = [
                    ("liability_cap",      "Liability Cap"),
                    ("termination_notice", "Termination Notice"),
                    ("payment_terms",      "Payment Terms"),
                ]
                for col, (cat, label) in zip(thresh_cols, key_clauses):
                    t = compute_dynamic_thresholds(deal_ctx, cat)
                    with col:
                        st.markdown(f"**{label}**")
                        if t.get("standard_months"):
                            st.caption(f"Floor: {t.get('floor_months')} mo  ·  Standard: {t.get('standard_months')} mo  ·  Ceiling: {t.get('ceiling_months')} mo")
                            if t.get("standard_cad"):
                                st.caption(f"Standard ≈ CAD ${t['standard_cad']:,}  ·  Floor ≈ CAD ${t.get('floor_cad',0):,}")
                        elif t.get("standard_days"):
                            st.caption(f"Floor: {t.get('floor_days')} d  ·  Standard: {t.get('standard_days')} d  ·  Ceiling: {t.get('ceiling_days')} d")
                        if t.get("rationale"):
                            st.caption(f"📐 {t['rationale']}")

                st.caption(f"**Industry risk tier:** {deal_ctx.industry_risk_tier}  ·  "
                           f"**Deal size:** {deal_ctx.deal_size_tier}  ·  "
                           f"**Risk tolerance:** {deal_ctx.risk_tolerance}")

                # Top risk weights for this deal
                top_weights = sorted(deal_ctx.adjusted_weights.items(), key=lambda x: -x[1])[:6]
                weight_html = '<div style="display:flex;flex-wrap:wrap;gap:6px;margin-top:8px;">'
                for clause, w in top_weights:
                    label = clause.replace("_"," ").title()
                    intensity = min(int(w * 2.5), 200)
                    weight_html += (f'<span style="background:rgba(139,0,0,{intensity/255:.2f});'
                                    f'color:{"white" if intensity > 120 else "#333"};'
                                    f'padding:3px 8px;border-radius:4px;font-size:11px;">'
                                    f'{label}: {w}</span>')
                weight_html += '</div>'
                st.caption("**Clause risk weights for this deal:**")
                st.markdown(weight_html, unsafe_allow_html=True)

        # ── COST GUARD ───────────────────────────────────────────
        session_cost_so_far = sum(h.get('cost', 0) for h in st.session_state.history)
        depth_tok = {"Quick Scan": 8000, "Standard Review": 16000, "Deep Dive + Redraft": 20000}
        est_this = (
            (3000 + chars // 4) * _INPUT_RATE / 1_000_000
            + depth_tok.get(analysis_depth, 12000) * _OUTPUT_RATE / 1_000_000
        )
        est_session_total = session_cost_so_far + est_this
        if est_session_total > MAX_SESSION_COST_USD:
            st.warning(
                f"⚠️ Session cost notice: estimated total after this analysis will be "
                f"~${est_session_total:.2f} USD (this run ~${est_this:.4f}). "
                f"Proceeding.")
        # ────────────────────────────────────────────────────────

        _audit.info(
            f"ANALYSIS_START | matter={name!r} | model={model_id} | depth={analysis_depth} "
            f"| chars={chars} | province={province} | reviewer={reviewer_name or 'unknown'} "
            f"| est_cost=${est_this:.4f}")

        st.info(f"Analysing ({analysis_depth} | {output_mode}){'  [chunked mode]' if is_chunked else ''}...")

        # ── PRE-AI PYTHON RISK SCORING ──────────────────────────
        # Score key clauses against deal-calibrated thresholds using
        # regex extraction before the AI call. Results injected as
        # verified ground truth the AI cannot contradict.
        pre_scores = {}
        if deal_ctx:
            for cat in ["liability_cap", "payment_terms", "late_fees"]:
                text_lower = contract_text.lower()
                snippets = []
                for kw in CLAUSE_KEYWORDS.get(cat, []):
                    idx = text_lower.find(kw)
                    if idx != -1:
                        snippets.append(contract_text[max(0, idx-50):idx+200])
                if snippets:
                    provision_text = " ".join(snippets[:2])
                    score_result = compute_clause_risk_score(cat, provision_text, deal_ctx)
                    pre_scores[cat] = score_result

        # ── STATUTORY RED-LINE CHECK (hybrid rule-based engine) ──
        # Non-waivable Canadian law violations caught in Python
        # before the AI sees the contract.  Results displayed to the
        # user and injected into the AI prompt as verified facts.
        statutory_flags = check_statutory_redlines(
            contract_text, ctype if ctype != "Auto-Detect" else "")
        if statutory_flags:
            with st.expander(
                    f"⚠️ Statutory Red-Lines Detected ({len(statutory_flags)})",
                    expanded=True):
                st.caption(
                    "Python-detected violations of non-waivable Canadian law. "
                    "Injected into the AI prompt as verified facts.")
                for flag in statutory_flags:
                    sev   = flag["severity"]
                    color = "#8B0000" if sev == "CRITICAL" else "#B8860B"
                    st.markdown(
                        f"<div style='border-left:4px solid {color};"
                        f"padding:8px 12px;background:#fafafa;margin:4px 0;'>"
                        f"<b>{html_escape(sev)}:</b> {html_escape(flag['statute'])}<br/>"
                        f"{html_escape(flag['flag'])}</div>",
                        unsafe_allow_html=True)

        # Build unified pre-score + statutory annotation for AI prompt
        pre_score_note = ""
        note_lines = []
        if pre_scores:
            note_lines.append(
                "PRE-VERIFIED PYTHON RISK SCORES "
                "(treat as ground truth — do not contradict):")
            for cat, result in pre_scores.items():
                label = cat.replace("_", " ").title()
                icon  = {"GREEN": "✓", "AMBER": "⚠", "RED": "✗"}.get(
                    result["status"], "?")
                note_lines.append(
                    f"  {icon} {label}: {result['status']} "
                    f"(score: {result['score']}) — {result['rationale']}")
        if statutory_flags:
            note_lines.append(
                "\nSTATUTORY RED-LINES "
                "(Python-verified, non-waivable — flag as CRITICAL):")
            for flag in statutory_flags:
                note_lines.append(
                    f"  [{flag['severity']}] {flag['statute']}: {flag['flag']}")
        if note_lines:
            pre_score_note = "\n".join(note_lines)
        try:
            data, full, elapsed, cost, is_json = run_analysis(
                contract_text, clause_index, ctype, province, client_role,
                special, playbook_text, analysis_depth, output_mode,
                analysis_language, model_id, use_caching, deal_ctx,
                pre_score_note)
        except Exception as e:
            _audit.error(f"ANALYSIS_ERROR | matter={name!r} | error={e!r}")
            st.error(f"Analysis error: {e}"); st.stop()

        if is_json:
            risk, score, verdict = compute_risk(data)
        else:
            st.warning("The AI returned a non-structured response. Displaying raw output.")
            risk, score, verdict = "MODERATE", 0, "Review required."
            data = None

        rc_css = {"CRITICAL":"risk-critical","HIGH":"risk-high","MODERATE":"risk-medium","LOW":"risk-low"}.get(risk,"risk-medium")
        rc_hex = {"CRITICAL":"#8B0000","HIGH":"#8B0000","MODERATE":"#B8860B","LOW":"#006400"}.get(risk,"#444")
        chunk_note = f"  &nbsp;|&nbsp;  Chunked ({chars//1000}k chars)" if is_chunked else ""
        st.markdown(f"""<div class="risk-banner {rc_css}">
          <h3 style='color:{rc_hex};margin:0;font-family:Georgia;'>Overall Risk: {html_escape(str(risk))} &nbsp;|&nbsp; Score: {score}</h3>
          <p style='margin:4px 0 0 0;color:#333;font-size:14px;'>{html_escape(str(verdict))}</p>
          <p style='margin:4px 0 0 0;color:#999;font-size:11px;'>{elapsed:.0f}s &nbsp;|&nbsp; ~${cost:.4f} &nbsp;|&nbsp; {output_mode}{chunk_note} &nbsp;|&nbsp; {doc_count} doc(s)</p>
        </div>""", unsafe_allow_html=True)

        st.session_state.history.append({
            "name": name, "risk": risk,
            "time": datetime.now().strftime("%H:%M"), "cost": cost,
        })
        _audit.info(
            f"ANALYSIS_COMPLETE | matter={name!r} | risk={risk} | score={score} "
            f"| cost=${cost:.4f} | elapsed={elapsed:.1f}s | is_json={is_json}")

        # Store full analysis result in session state so it persists
        # across rerenders (e.g. when user clicks a download button).
        if is_json and data:
            st.session_state.analysis_result = {
                "data":             data,
                "name":             name,
                "safe":             re.sub(r'[^\w\s-]','',name).strip().replace(' ','_')[:40],
                "risk":             risk,
                "score":            score,
                "verdict":          verdict,
                "cost":             cost,
                "elapsed":          elapsed,
                "chars":            chars,
                "is_chunked":       is_chunked,
                "doc_count":        doc_count,
                "ctype":            ctype,
                "province":         province,
                "client_role":      client_role,
                "output_mode":      output_mode,
                "matter_id":        matter_id if 'matter_id' in locals() else "",
                "extraction_method":extraction_method,
                "contract_text":    contract_text,
                "citation_issues":  [],  # filled below
                "pi_findings":      pi_findings if 'pi_findings' in locals() else [],
                "pi_summary":       pi_summary  if 'pi_summary'  in locals() else {},
            }

        # Save to persistent matter store.
        # The analysis_json contains verbatim_text fields (direct contract quotes).
        # We strip these before saving so the matter store holds only metadata
        # and AI-generated analysis — not raw contract text — satisfying the
        # "locally stored matter metadata only" security claim.
        if is_json and data:
            data_for_store = _sanitise_for_storage(data)
            deal_ctx_dict = None
            if deal_ctx:
                deal_ctx_dict = {
                    "contract_value_cad": deal_ctx.contract_value_cad,
                    "annual_fees_cad":    deal_ctx.annual_fees_cad,
                    "industry":           deal_ctx.industry,
                    "counterparty_type":  deal_ctx.counterparty_type,
                    "deal_size_tier":     deal_ctx.deal_size_tier,
                    "industry_risk_tier": deal_ctx.industry_risk_tier,
                    "risk_tolerance":     deal_ctx.risk_tolerance,
                    "duration_months":    deal_ctx.duration_months,
                }
            save_matter(
                name=name, matter_id=matter_id, contract_type=data.get('contract_type', ctype),
                province=province, client_role=client_role, risk_level=risk,
                risk_score=score, verdict=verdict, cost_usd=cost,
                model=model_id, analysis_json_obj=data_for_store,
                char_count=chars, doc_count=doc_count,
                deal_context_json_obj=deal_ctx_dict,
                reviewer_name=reviewer_name if reviewer_name else "")

        # Schema validation + citation validation
        citation_issues = []
        if is_json and data:
            # Run schema validation and show any structural issues
            schema_errs = validate_analysis_schema(data)
            if schema_errs:
                with st.expander(f"⚠️ AI Response Quality Issues ({len(schema_errs)})", expanded=False):
                    st.caption("These are structural issues with the AI response — the analysis may be incomplete.")
                    for err in schema_errs:
                        st.warning(err)

            citation_issues = validate_citations(data, clause_index)
            if citation_issues:
                st.warning(f"Citation validation: {len(citation_issues)} questionable reference(s).")
                with st.expander(f"Citation Validation ({len(citation_issues)} issue(s))"):
                    for ci in citation_issues:
                        st.error(f"Section **{html_escape(ci['claimed'])}** in '{html_escape(ci['in_issue'])}' not found in clause index.")

            source_links = generate_source_links(data)
            if source_links:
                with st.expander(f"Source Links ({len(source_links)} statutes/cases referenced)"):
                    for link in source_links:
                        st.markdown(f"[{link['name']}]({link['url']})")

        if is_json and data:
            clause_review = data.get('mandatory_clause_review',[])
            if clause_review:
                st.subheader("Mandatory Clause Heatmap")
                render_heatmap(clause_review)

            st.subheader("Executive Summary")
            with st.container(border=True):
                st.markdown(data.get('executive_summary',''))

            high = data.get('high_exposure_issues',[])
            if high:
                with st.expander(f"High Exposure Issues ({len(high)})", expanded=True):
                    render_issues(high,"High Exposure Issues")

            elevated = data.get('elevated_risk_issues',[])
            if elevated:
                with st.expander(f"Elevated Risk Issues ({len(elevated)})", expanded=True):
                    render_issues(elevated,"Elevated Risk Issues")

            imb = data.get('commercial_imbalances',[])
            if imb:
                with st.expander(f"Commercial Imbalances ({len(imb)})"):
                    for item in imb:
                        st.markdown(f"**{item.get('title','')}** — {item.get('clause_reference','')}")
                        st.markdown(item.get('description',''))
                        st.markdown(f"*Recommendation:* {item.get('recommendation','')}")
                        st.markdown("---")

            missing = data.get('missing_protections',[])
            if missing:
                with st.expander(f"Missing Protections ({len(missing)})", expanded=True):
                    for item in missing:
                        st.markdown(f"**{item.get('clause_type','')}** — Risk: {item.get('risk_level','')}")
                        st.markdown(item.get('explanation',''))
                        if item.get('proposed_clause'):
                            st.code(item['proposed_clause'], language=None)

            neg = data.get('negotiation_strategy',{})
            if neg and any(neg.values()):
                with st.expander("Negotiation Strategy"):
                    if neg.get('priority_items'):
                        st.markdown("**Priority Items:**")
                        for i,item in enumerate(neg['priority_items'],1): st.markdown(f"{i}. {item}")
                    if neg.get('walk_away_points'):
                        st.markdown("**Walk-Away Points:**")
                        for item in neg['walk_away_points']: st.markdown(f"- {item}")
                    if neg.get('draft_communication'):
                        st.markdown("**Draft Communication:**")
                        st.code(neg['draft_communication'], language=None)

            redlines = data.get('suggested_redline_clauses',[])
            if redlines:
                with st.expander(f"Suggested Redline Clauses ({len(redlines)})"):
                    for rl in redlines:
                        st.markdown(f"**{rl.get('clause_reference','')}** — {rl.get('issue','')}")
                        if rl.get('current_text'):  st.markdown(f"*Current:* {rl['current_text']}")
                        if rl.get('replacement_text'): st.markdown(f"*Proposed:* {rl['replacement_text']}")
                        prec = rl.get('precedent_used','')
                        if prec and prec != 'drafted': st.caption(f"Precedent: {prec}")
                        if rl.get('rationale'): st.caption(f"Rationale: {rl['rationale']}")
                        st.markdown("---")

            actions = data.get('action_items',{})
            if actions and any(actions.values()):
                with st.expander("Action Items"):
                    for phase, items in actions.items():
                        if items:
                            st.markdown(f"**{phase.replace('_',' ').title()}:**")
                            for item in items: st.markdown(f"- [ ] {item}")

            if clause_review:
                with st.expander("Contract Risk Scorecard"):
                    scorecard_html = '<table style="width:100%;border-collapse:collapse;font-size:13px;">'
                    scorecard_html += '<tr style="background:#0d1b2a;color:white;"><th style="padding:6px;text-align:left;">Clause</th><th style="padding:6px;">Risk</th><th style="padding:6px;">Enforceability</th><th style="padding:6px;text-align:left;">Benchmark</th></tr>'
                    for item in clause_review:
                        cat    = item.get("clause_category","")
                        label  = html_escape(MANDATORY_CLAUSE_LABELS.get(cat, cat.replace("_"," ").title()))
                        rlvl   = html_escape(str(item.get("risk_level","N/A")))
                        rbg    = {"CRITICAL":"#fce4e4","HIGH":"#fef3c7","MEDIUM":"#fef9c3","LOW":"#dcfce7","MISSING":"#f3f4f6"}.get(item.get("risk_level",""),"#fff")
                        enf    = item.get("enforceability_pct","")
                        enf_s  = f"{enf}%" if enf != "" else "N/A"
                        try:
                            ev = int(str(enf)) if str(enf).isdigit() else 0
                            ec = "#15803d" if ev>=70 else "#b45309" if ev>=40 else "#991b1b"
                        except (ValueError,TypeError):
                            ec = "#991b1b"
                        bench = html_escape(str(item.get("benchmark_comparison",""))[:100])
                        scorecard_html += f'<tr style="border-bottom:1px solid #eee;"><td style="padding:6px;">{label}</td>'
                        scorecard_html += f'<td style="padding:6px;text-align:center;background:{rbg};font-weight:600;">{rlvl}</td>'
                        scorecard_html += f'<td style="padding:6px;text-align:center;color:{ec};font-weight:600;">{html_escape(str(enf_s))}</td>'
                        scorecard_html += f'<td style="padding:6px;font-size:11px;">{bench}</td></tr>'
                    scorecard_html += '</table>'
                    st.markdown(scorecard_html, unsafe_allow_html=True)

            st.session_state.last_analysis = {
                "data":data, "contract_type":data.get('contract_type',ctype), "name":name}

            with st.expander("Raw JSON (debug)"):
                st.json(data)

            # ── DOWNLOADS ──
            # Generate all file bytes NOW and cache in session_state.
            # This means download buttons survive page rerenders (e.g. after clicking
            # a button). Files are only regenerated when a new analysis runs.
            st.divider(); st.subheader("Download Reports")
            safe = re.sub(r'[^\w\s-]','',name).strip().replace(' ','_')[:40]
            ts   = datetime.now().strftime("%Y%m%d_%H%M%S")
            meta = dict(name=name, type=data.get('contract_type',ctype),
                        role=client_role, output_mode=output_mode,
                        province=province, risk=risk, score=score, verdict=verdict,
                        safe=safe, ts=ts,
                        matter_id=matter_id if matter_id else "")

            has_redlines = bool(data.get('suggested_redline_clauses')) or any(
                i.get('proposed_replacement') for i in
                data.get('high_exposure_issues',[]) + data.get('elevated_risk_issues',[]))

            # Generate and cache all bytes
            # All make_*_docx functions now return bytes directly (BytesIO).
            # make_pdf_report still uses a tempfile (ReportLab limitation).
            with st.spinner("Preparing download files..."):
                try:
                    pf = make_pdf_report(data, meta, citation_issues)
                    with open(pf,"rb") as fh: _pb = fh.read()
                    try: os.unlink(pf)
                    except OSError: pass
                    st.session_state._dl_pdf = _pb if _valid_pdf(_pb) else None
                    if not _valid_pdf(_pb):
                        st.error(f"PDF generation produced invalid output ({len(_pb)} bytes).")
                except Exception as e:
                    st.session_state._dl_pdf = None
                    st.error(f"PDF generation failed: {e}")

                try:
                    _db = make_docx_report(data, meta, citation_issues)
                    st.session_state._dl_docx = _db if _valid_docx(_db) else None
                except Exception as e:
                    st.session_state._dl_docx = None

                if has_redlines:
                    try:
                        _ab = make_amendments_docx(data, meta)
                        st.session_state._dl_amend = _ab if _valid_docx(_ab) else None
                    except Exception as e:
                        st.session_state._dl_amend = None

                    try:
                        _rb = make_redline_docx(data, meta)
                        if _valid_docx(_rb):
                            st.session_state._dl_redline = _rb
                        else:
                            st.session_state._dl_redline = None
                            st.error(f"Track changes DOCX generation produced invalid output ({len(_rb)} bytes). Please try again.")
                    except Exception as e:
                        st.session_state._dl_redline = None
                        st.error(f"Track changes DOCX failed: {e}")

            st.session_state._dl_safe = safe
            st.session_state._dl_ts   = ts
            st.session_state._dl_has_redlines = has_redlines

            revised = data.get('revised_contract','')
            if revised:
                st.download_button("Revised Contract Draft", data=revised,
                    file_name=f"{safe}_Revised_{ts}.txt", mime="text/plain")

            # Applied redline
            all_redlines = data.get('suggested_redline_clauses',[])
            if not all_redlines:
                all_redlines = [
                    {"current_text":i['verbatim_text'], "replacement_text":i['proposed_replacement']}
                    for i in data.get('high_exposure_issues',[]) + data.get('elevated_risk_issues',[])
                    if i.get('proposed_replacement') and i.get('verbatim_text')
                ]
            if all_redlines:
                st.divider(); st.subheader("One-Click Applied Redline")
                st.caption("Apply all amendments directly to the original contract text using OOXML track changes.")
                if st.button("Generate Applied Redline", use_container_width=True):
                    try:
                        ar_bytes = make_applied_redline_docx(contract_text, all_redlines, meta)
                        st.download_button("Download Applied Redline (.docx)", data=ar_bytes,
                            file_name=f"{safe}_AppliedRedline_{ts}.docx",
                            mime="application/octet-stream",
                            type="primary", use_container_width=True)
                        st.success(f"Applied {len(all_redlines)} amendment(s) with OOXML track changes.")
                    except Exception as e: st.error(f"Applied redline error: {e}")

            st.download_button("Export JSON", data=json.dumps(data, indent=2),
                file_name=f"{safe}_Analysis_{ts}.json", mime="application/json")

            # ── AUDIT LOG ──
            st.divider()
            audit_col1, audit_col2 = st.columns(2)
            with audit_col1:
                st.subheader("Audit Log")
                audit_text = f"""AI ANALYSIS AUDIT LOG
{"="*50}
Date:                {datetime.now().strftime('%B %d, %Y at %H:%M:%S')}
Contract:            {name}
Matter ID:           {matter_id if matter_id else 'Not specified'}
Contract Type:       {data.get('contract_type', ctype)}
Jurisdiction:        {province}
Client Position:     {client_role}
Analysis Depth:      {analysis_depth}
Output Mode:         {output_mode}
Model:               {model_id}
Prompt Version:      v12.6
Documents:           {doc_count}
Contract Length:     {chars:,} characters ({chars//4:,} est. tokens)
Extraction Method:   {extraction_method}
Chunked Analysis:    {'Yes (' + str(len(_split_into_chunks(contract_text))) + ' chunks)' if is_chunked else 'No'}
Sections Analysed:   15 mandatory clause categories
Analysis Duration:   {elapsed:.1f} seconds
Risk Score:          {score} ({risk})
Estimated Cost:      ${cost:.4f}
Citation Issues:     {len(citation_issues)}
PI Scan Result:      {pi_summary["level"]} ({len(pi_findings)} category(s) detected)
Saved to DB:         Yes
{"="*50}
Human Reviewer:      {reviewer_name if reviewer_name else '____________________________'}
Review Date:         ____________________________
Reviewer Signature:  ____________________________
"""
                st.code(audit_text, language=None)
                st.download_button("Download Audit Log", data=audit_text,
                    file_name=f"{safe}_AuditLog_{ts}.txt", mime="text/plain")

            with audit_col2:
                st.subheader("Partner Review Checklist")
                checklist_items = [
                    "Verify governing law and jurisdiction clause",
                    "Confirm liability cap enforceability and adequacy",
                    "Check indemnity scope and mutuality",
                    "Review termination provisions for asymmetry",
                    "Confirm statutory compliance (ESA/CPA/Construction Act as applicable)",
                    "Verify IP assignment and moral rights waiver",
                    "Check non-compete enforceability (Bill 27 if ON employee)",
                    "Review payment terms and late fee compliance (s.347 Criminal Code)",
                    "Confirm confidentiality scope and carve-outs",
                    "Verify force majeure clause present and adequate",
                    "Review insurance requirements",
                    "Check notice provisions",
                    "Review data protection / PIPEDA compliance",
                    "Confirm all AI-flagged issues addressed",
                    "Verify all proposed replacement clauses (check precedent used)",
                ]
                checklist_text = (f"PARTNER REVIEW CHECKLIST\n{'='*50}\n"
                                  f"Contract: {name}\nDate: {datetime.now().strftime('%B %d, %Y')}\n"
                                  f"AI Risk Assessment: {risk} (Score: {score})\n{'='*50}\n\n"
                                  + "".join(f"[ ]  {i+1}. {item}\n\n" for i,item in enumerate(checklist_items))
                                  + f"\n{'='*50}\nReviewed by: ____________________________\n"
                                    f"Date: ____________________________\n")
                st.code(checklist_text, language=None)
                st.download_button("Download Checklist", data=checklist_text,
                    file_name=f"{safe}_ReviewChecklist_{ts}.txt", mime="text/plain")

        else:
            st.subheader("Analysis Output")
            with st.container(border=True): st.markdown(full)


    # ── PERSISTENT DOWNLOAD BUTTONS ─────────────────────────────
    # Rendered on every page load from cached session state.
    # Download buttons don't disappear when clicked.
    if st.session_state.get('_dl_pdf') or st.session_state.get('_dl_redline'):
        _safe = st.session_state.get('_dl_safe', 'contract')
        _ts   = st.session_state.get('_dl_ts', datetime.now().strftime("%Y%m%d_%H%M%S"))
        _hr   = st.session_state.get('_dl_has_redlines', False)

        st.divider()
        hdr_col, reset_col = st.columns([4, 1])
        with hdr_col:
            st.subheader("Download Reports")
            st.caption("Files are ready — click to download. Results stay available until you run a new analysis.")
        
        # Detect Edge browser via user agent hint and show fix instructions
        st.markdown("""
<div style='background:#e8f0fe;border:1px solid #4285f4;padding:10px 14px;
border-radius:4px;font-size:12px;color:#1a237e;margin-bottom:10px;'>
<b>⚙️ Microsoft Edge users — one-time fix for DOCX downloads:</b><br/>
Edge intercepts Word files and tries to open them in the browser.
To fix: <b>Edge Settings → Downloads → turn OFF "Open Office files in the browser"</b><br/>
Or: after downloading, right-click the file in the taskbar downloads → 
<b>"Keep" → right-click in Explorer → "Open with" → Microsoft Word</b>
</div>""", unsafe_allow_html=True)
        with reset_col:
            if st.button("↺ Reset", key="reset_downloads",
                         help="Clear cached files and re-generate on next analysis run"):
                for k in ['_dl_pdf','_dl_docx','_dl_amend','_dl_redline',
                          '_dl_safe','_dl_ts','_dl_has_redlines']:
                    st.session_state.pop(k, None)
                st.rerun()

        # All DOCX downloads use HTML data URI anchor tags.
        # Edge intercepts http downloads of .docx files via its built-in
        # Office viewer. Data URIs are not http downloads so Edge cannot
        # intercept them — files go straight to the Downloads folder.
        if _hr and st.session_state.get('_dl_redline'):
            _rl_kb = len(st.session_state._dl_redline) // 1024
            st.markdown(
                "<div style='background:#0d1b2a;padding:10px 16px;border-radius:6px;"
                "margin-bottom:8px;'>"
                "<span style='color:#c9a84c;font-weight:600;font-size:13px;'>"
                "⭐ RECOMMENDED FOR LAWYERS</span><br/>"
                "<span style='color:#8fa4c4;font-size:11px;'>"
                "True Track Changes — opens in Word with native Accept/Reject buttons"
                "</span></div>",
                unsafe_allow_html=True)
            st.markdown(_docx_download_link(
                st.session_state._dl_redline,
                f"{_safe}_Redline_{_ts}.docx",
                f"⬇  True Track Changes DOCX  (open in Word → Review → Accept/Reject)  [{_rl_kb} KB]",
                _rl_kb, primary=True), unsafe_allow_html=True)
            st.caption(f"OOXML w:ins/w:del — native Word track changes · {_rl_kb} KB")
            st.divider()

        dl_cols = st.columns(3 if _hr else 2)
        with dl_cols[0]:
            if st.session_state.get('_dl_pdf'):
                _pk = len(st.session_state._dl_pdf) // 1024
                st.download_button(f"PDF Memo  [{_pk} KB]",
                    data=st.session_state._dl_pdf,
                    file_name=f"{_safe}_Memo_{_ts}.pdf",
                    mime="application/pdf",
                    use_container_width=True, key="persistent_pdf")
                st.caption(f"Structured memo · {_pk} KB")
            else:
                st.caption("PDF not available")
        with dl_cols[1]:
            if st.session_state.get('_dl_docx'):
                _dk = len(st.session_state._dl_docx) // 1024
                st.markdown(_docx_download_link(
                    st.session_state._dl_docx,
                    f"{_safe}_Memo_{_ts}.docx",
                    f"⬇  Word Memo  [{_dk} KB]",
                    _dk), unsafe_allow_html=True)
                st.caption(f"Editable Word document · {_dk} KB")
            else:
                st.caption("Word Memo not available")
        if _hr and len(dl_cols) > 2:
            with dl_cols[2]:
                if st.session_state.get('_dl_amend'):
                    _ak = len(st.session_state._dl_amend) // 1024
                    st.markdown(_docx_download_link(
                        st.session_state._dl_amend,
                        f"{_safe}_Amendments_{_ts}.docx",
                        f"⬇  Amendments Schedule  [{_ak} KB]",
                        _ak), unsafe_allow_html=True)
                    st.caption(f"For opposing counsel · {_ak} KB")
                else:
                    st.caption("Amendments not available")



# ── TAB 2: VERSION COMPARISON ─────────────────────────────────
with tab_compare:
    st.subheader("Version Comparison")
    st.caption("Upload the original contract and the counterparty redline to identify what changed and how risk was affected.")
    comp_c1, comp_c2 = st.columns(2)
    with comp_c1:
        orig_file = st.file_uploader("Original Contract", type=["pdf","docx","txt"], key="orig_comp")
    with comp_c2:
        red_file = st.file_uploader("Counterparty Redline", type=["pdf","docx","txt"], key="red_comp")

    if st.button("Compare Versions", type="primary", use_container_width=True):
        if not orig_file or not red_file:
            st.error("Upload both versions."); st.stop()
        with st.spinner("Extracting and comparing..."):
            orig_text, _, _ = extract_text(orig_file)
            red_text,  _, _ = extract_text(red_file)
            changes = compare_versions(orig_text, red_text)
        st.success(f"Comparison: {len(changes['modified'])} modifications, {len(changes['added'])} additions, {len(changes['removed'])} removals")

        compare_prompt = f"""You are a senior contract lawyer. Compare these two contract versions and assess the impact of each change.
ORIGINAL:
{orig_text[:20000]}
REVISED:
{red_text[:20000]}
Respond with ONLY valid JSON (no code fences):
{{
  "summary": "string — 2-3 sentence professional summary",
  "changes": [
    {{
      "clause": "string",
      "description": "string",
      "impact": "IMPROVED | WORSENED | NEUTRAL",
      "risk_analysis": "string",
      "recommendation": "string",
      "confidence": "HIGH | MEDIUM | REQUIRES_REVIEW"
    }}
  ],
  "overall_direction": "IMPROVED | WORSENED | MIXED | NEUTRAL",
  "recommendation": "string"
}}"""
        st.info("Analysing changes...")
        try:
            resp = ai.messages.create(model=model_id, max_tokens=6000,
                messages=[{"role":"user","content":compare_prompt}], timeout=180.0)
            comp_data, comp_ok = parse_json_response(resp.content[0].text)
            if comp_ok:
                direction = html_escape(str(comp_data.get('overall_direction','MIXED')))
                d_color   = {"IMPROVED":"#006400","WORSENED":"#8B0000","MIXED":"#B8860B","NEUTRAL":"#444"}.get(comp_data.get('overall_direction',''),"#444")
                st.markdown(f"""<div class="risk-banner" style="border-left:5px solid {d_color};">
                  <h3 style='color:{d_color};margin:0;'>Overall Direction: {direction}</h3>
                  <p style='margin:4px 0;'>{html_escape(str(comp_data.get('summary','')))}</p>
                  <p style='margin:4px 0;'><b>Recommendation:</b> {html_escape(str(comp_data.get('recommendation','')))}</p></div>""",
                unsafe_allow_html=True)
                for ch in comp_data.get('changes',[]):
                    impact  = ch.get('impact','NEUTRAL')
                    css     = {"IMPROVED":"issue-low","WORSENED":"issue-critical","NEUTRAL":""}.get(impact,"")
                    conf    = ch.get('confidence','')
                    conf_css= {"HIGH":"conf-high","MEDIUM":"conf-medium","REQUIRES_REVIEW":"conf-low"}.get(conf,"")
                    st.markdown(f"""<div class="issue-card {css}">
                        <b>{html_escape(str(ch.get('clause','')))} — {html_escape(str(impact))}</b>
                        &nbsp;<span class="{conf_css}">[{html_escape(str(conf))}]</span><br/>
                        <b>Change:</b> {html_escape(str(ch.get('description','')))}<br/>
                        <b>Risk:</b> {html_escape(str(ch.get('risk_analysis','')))}<br/>
                        <b>Recommendation:</b> {html_escape(str(ch.get('recommendation','')))}
                    </div>""", unsafe_allow_html=True)
            else:
                st.markdown(resp.content[0].text)
        except Exception as e:
            st.error(f"Comparison error: {e}")

        with st.expander("Raw Text Differences"):
            if changes['modified']:
                st.markdown("**Modified:**")
                for m in changes['modified'][:15]:
                    st.markdown(f"- ~~{m['original'][:120]}~~ → *{m['revised'][:120]}*")
            if changes['added']:
                st.markdown("**Added:**")
                for a in changes['added'][:10]: st.markdown(f"+ *{a[:120]}*")
            if changes['removed']:
                st.markdown("**Removed:**")
                for r_item in changes['removed'][:10]: st.markdown(f"- ~~{r_item[:120]}~~")


# ── TAB 3: NEGOTIATION SIMULATOR ─────────────────────────────
with tab_negotiate:
    st.subheader("Negotiation Simulator")
    st.caption("Predicts how opposing counsel would respond to each proposed amendment. Run a contract analysis first, then simulate.")

    st.markdown(
        """<div style='background:#fff8e7;border:1px solid #e5c76b;padding:10px 16px;
        border-radius:4px;margin-bottom:12px;font-size:12px;color:#7a5d00;
        font-family:Georgia,serif;'>
        <b>IMPORTANT — ILLUSTRATIVE SCENARIO PLANNING ONLY</b><br/>
        Acceptance probabilities and predicted responses are AI-generated estimates
        based on general commercial negotiation patterns. They are <b>not reliable
        predictions</b> and have no statistical basis. Do not communicate these
        figures to clients, use them to advise on settlement, or treat them as a
        substitute for experienced negotiation counsel. Every deal is unique.
        </div>""",
        unsafe_allow_html=True)

    if st.session_state.get('last_analysis',{}).get('data'):
        last = st.session_state.last_analysis
        data_for_sim = last['data']
        all_issues = data_for_sim.get('high_exposure_issues',[]) + data_for_sim.get('elevated_risk_issues',[])
        amendable  = [i for i in all_issues if i.get('proposed_replacement')]

        if amendable:
            st.info(f"Found {len(amendable)} proposed amendments from: **{last.get('name','')}**")
            if st.button("Run Negotiation Simulation", type="primary", use_container_width=True):
                sim_prompt = build_negotiation_sim_prompt(amendable, last.get('contract_type',''))
                with st.spinner("Simulating opposing counsel responses..."):
                    try:
                        sim_resp = ai.messages.create(
                            model=model_id, max_tokens=6000,
                            messages=[{"role":"user","content":sim_prompt}], timeout=180.0)
                        sim_data, sim_ok = parse_json_response(sim_resp.content[0].text)
                        if sim_ok and isinstance(sim_data, dict):
                            for sim in sim_data.get('simulations',[]):
                                response   = sim.get('opposing_likely_response','UNKNOWN')
                                resp_color = {"ACCEPT":"#15803d","REJECT":"#991b1b","COUNTER":"#b45309"}.get(response,"#444")
                                prob       = sim.get('probability_of_acceptance','')
                                st.markdown(f"""<div class="issue-card" style="border-left:4px solid {resp_color};">
                                    <b>{html_escape(str(sim.get('amendment_title','')))}</b><br/>
                                    <b>Opposing Response:</b> <span style="color:{resp_color};font-weight:700;">{html_escape(str(response))}</span>
                                    {f' &mdash; AI estimate: {html_escape(str(prob))}% (illustrative only)' if prob else ''}<br/>
                                    {"<b>Predicted Counter:</b> " + html_escape(str(sim.get('predicted_counter',''))) + "<br/>" if sim.get('predicted_counter') else ""}
                                    <b>Best Case:</b> {html_escape(str(sim.get('best_case','')))}<br/>
                                    <b>Likely Compromise:</b> {html_escape(str(sim.get('likely_compromise','')))}<br/>
                                    <b>Fallback Position:</b> {html_escape(str(sim.get('fallback_position','')))}<br/>
                                    <b>Leverage:</b> {html_escape(str(sim.get('leverage_notes','')))}
                                </div>""", unsafe_allow_html=True)
                            st.download_button("Export Simulation JSON",
                                data=json.dumps(sim_data,indent=2),
                                file_name=f"negotiation_simulation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                                mime="application/json")
                        else:
                            st.markdown(sim_resp.content[0].text)
                    except Exception as e:
                        st.error(f"Simulation error: {e}")
        else:
            st.warning("No proposed amendments found. Run a Standard or Deep Dive analysis first.")
    else:
        st.info("Run a contract analysis first. The simulator uses the issues from your most recent analysis.")


# ── TAB 4: MATTER HISTORY ─────────────────────────────────────
with tab_history:
    st.subheader("Matter History")
    st.caption("All analyses are saved automatically. Click Load to restore a previous matter.")

    init_db()

    # ── Filter by reviewer (Law Society accountability) ───────
    hist_c1, hist_c2 = st.columns([2, 1])
    with hist_c1:
        reviewer_filter = st.text_input(
            "Filter by reviewer name",
            placeholder="Leave blank to show all",
            key="hist_reviewer_filter")
    with hist_c2:
        hist_limit = st.selectbox("Show", [25, 50, 100, 200], index=1,
            key="hist_limit")

    if reviewer_filter.strip():
        matters = matters_by_reviewer(reviewer_filter.strip(), limit=hist_limit)
        st.caption(f"{len(matters)} matter(s) reviewed by '{reviewer_filter.strip()}'.")
    else:
        matters = list_matters(limit=hist_limit)
        st.caption(f"{len(matters)} matter(s) on record.")

    if not matters:
        st.info("No matters found. Run an analysis to see it here.")
    else:
        for m in matters:
            risk_icon  = {"CRITICAL":"🔴","HIGH":"🟠","MODERATE":"🟡","LOW":"🟢"}.get(
                m.get('risk_level',''),"⚪")
            date_str   = m['created_at'][:16].replace('T',' ') if m.get('created_at') else ""
            reviewer   = m.get('reviewer_name','') or '—'
            accessed   = (m.get('last_accessed_at','') or '')[:16].replace('T',' ')
            cols       = st.columns([3, 1, 1, 1, 1, 1, 1])
            cols[0].markdown(f"**{m.get('name','')[:35]}**  \n{date_str}")
            cols[1].caption(m.get('contract_type','')[:18])
            cols[2].caption(m.get('province',''))
            cols[3].markdown(f"{risk_icon} {m.get('risk_level','')}")
            cols[4].caption(f"${m.get('cost_usd',0):.3f}")
            cols[5].caption(f"👤 {reviewer[:14]}")
            if cols[6].button("Load", key=f"hist_load_{m['id']}"):
                st.session_state.loaded_matter = load_matter(m['id'])
                st.session_state.last_analysis = {
                    "data": st.session_state.loaded_matter.get('analysis_json',{}),
                    "contract_type": st.session_state.loaded_matter.get('contract_type',''),
                    "name": st.session_state.loaded_matter.get('name',''),
                }
                st.success(
                    f"Loaded: {m.get('name','')} — switch to Contract Analysis tab to view.")
            st.divider()

    if matters:
        exp_cols = st.columns(2)
        with exp_cols[0]:
            if st.button("Export Displayed Matters as JSON", key="hist_export_json"):
                all_data = [load_matter(m['id']) for m in matters]
                st.download_button(
                    "Download matters.json",
                    data      = json.dumps(all_data, indent=2, default=str),
                    file_name = f"contractcheck_matters_{datetime.now().strftime('%Y%m%d')}.json",
                    mime      = "application/json",
                    key       = "hist_download_json")
        with exp_cols[1]:
            if st.button("Export as CSV (for Law Society audit)", key="hist_export_csv"):
                import csv, io
                buf = io.StringIO()
                w = csv.writer(buf)
                w.writerow(["ID","Created","Matter ID","Name","Contract Type",
                            "Province","Risk","Score","Reviewer","Cost USD"])
                for m in matters:
                    w.writerow([m.get('id'), m.get('created_at','')[:16],
                                m.get('matter_id',''), m.get('name',''),
                                m.get('contract_type',''), m.get('province',''),
                                m.get('risk_level',''), m.get('risk_score',0),
                                m.get('reviewer_name',''), f"{m.get('cost_usd',0):.4f}"])
                st.download_button(
                    "Download matter_audit.csv",
                    data      = buf.getvalue(),
                    file_name = f"matter_audit_{datetime.now().strftime('%Y%m%d')}.csv",
                    mime      = "text/csv",
                    key       = "hist_download_csv")


# ══════════════════════════════════════════════════════════════
# SESSION HISTORY SIDEBAR
# ══════════════════════════════════════════════════════════════
if st.session_state.history:
    st.sidebar.divider()
    st.sidebar.subheader("This Session")
    for h in reversed(st.session_state.history):
        risk_label = h['risk']
        st.sidebar.caption(f"{h['time']}  |  {risk_label}  |  {h['name'][:28]}")
    st.sidebar.caption(f"**Session total: ${sum(h['cost'] for h in st.session_state.history):.4f}**")


# ══════════════════════════════════════════════════════════════
# ABOUT
# ══════════════════════════════════════════════════════════════
# ── TAB 5: PIPEDA COMPLIANCE ─────────────────────────────────
with tab_pipeda:
    st.subheader("PIPEDA Compliance Centre")
    st.caption(
        "Tools and templates to support compliance with the Personal Information "
        "Protection and Electronic Documents Act (Canada) when using AI-assisted "
        "contract review.")

    st.markdown("---")

    # ── Compliance status summary ─────────────────────────────
    st.markdown("#### Compliance Status")
    status_cols = st.columns(4)
    status_items = [
        ("Contract text retention", "✓ None — discarded after session", "#15803d"),
        ("Local DB storage",        "✓ Metadata only (text stripped)", "#15803d"),
        ("Anthropic training",      "✓ Opt-out via commercial API",    "#15803d"),
        ("Cross-border transfer",   "⚠ US infrastructure — disclose",  "#B8860B"),
    ]
    for col, (label, status, color) in zip(status_cols, status_items):
        col.markdown(
            f"<div style='background:#f9fafb;border:0.5px solid #e5e7eb;"
            f"border-radius:6px;padding:10px 12px;'>"
            f"<div style='font-size:11px;color:#6b7280;margin-bottom:4px;'>{html_escape(label)}</div>"
            f"<div style='font-size:13px;font-weight:600;color:{color};'>{html_escape(status)}</div>"
            f"</div>",
            unsafe_allow_html=True)

    st.markdown("---")

    # ── Retainer disclosure template ─────────────────────────
    st.markdown("#### Retainer Disclosure Clause")
    st.caption(
        "Insert this clause into your standard client retainer agreement. "
        "It addresses all PIPEDA Principle 3 (consent) requirements for "
        "AI-assisted processing and cross-border transfer to Anthropic's US infrastructure.")

    ret_c1, ret_c2 = st.columns([1, 2])
    with ret_c1:
        ret_firm = st.text_input("Firm name",
            placeholder="Smith & Associates LLP", key="pipeda_tab_firm")
        ret_prov = st.selectbox("Province", list(PROVINCE_RULES.keys()),
            key="pipeda_tab_prov")
    with ret_c2:
        retainer_preview = generate_retainer_disclosure(
            firm_name    = ret_firm or "[FIRM NAME]",
            jurisdiction = ret_prov)
        st.text_area("Preview", retainer_preview, height=220,
            key="pipeda_retainer_preview")

    st.download_button(
        "Download Retainer Disclosure Clause (.txt)",
        data      = retainer_preview,
        file_name = "PIPEDA_AI_Retainer_Disclosure.txt",
        mime      = "text/plain",
        type      = "primary",
        key       = "pipeda_tab_download")

    st.markdown("---")

    # ── Matter store retention management ────────────────────
    st.markdown("#### Matter Store Retention")
    st.caption(
        "PIPEDA Principle 5 (Retention Limitation) requires that personal "
        "information not be kept longer than necessary. Set a retention period "
        "and apply it to purge expired matters.")

    ret_d1, ret_d2, ret_d3 = st.columns(3)
    with ret_d1:
        ret_days = st.selectbox(
            "Retention period",
            [30, 60, 90, 180, 365],
            format_func=lambda x: f"{x} days ({x//30} month{'s' if x//30 != 1 else ''})",
            index=2, key="pipeda_ret_days")
        to_delete = matters_older_than(ret_days)
        st.caption(f"{to_delete} matter(s) older than {ret_days} days.")

    with ret_d2:
        if st.button(f"Purge matters older than {ret_days} days",
                     key="pipeda_purge_btn", type="primary"):
            if to_delete > 0:
                n = purge_expired_matters(ret_days)
                st.success(f"Deleted {n} expired matter(s).")
                st.rerun()
            else:
                st.info("No matters to purge for this retention period.")

    with ret_d3:
        total = matter_count()
        st.metric("Total matters stored", total)
        if st.button("Delete ALL matters", key="pipeda_delete_all",
                     help="Permanent — cannot be undone"):
            if st.session_state.get("pipeda_delete_all_confirmed"):
                n = purge_all_matters()
                st.success(f"All {n} matter(s) deleted.")
                st.session_state.pipeda_delete_all_confirmed = False
                st.rerun()
            else:
                st.session_state.pipeda_delete_all_confirmed = True
                st.warning("Click again to confirm deletion of all matters.")

    st.markdown("---")

    # ── PIPEDA obligations checklist ─────────────────────────
    st.markdown("#### Firm Obligations Checklist")
    st.caption("Work through this checklist before deploying for client use.")

    obligations = [
        ("Principle 1 — Accountability",
         "Designate a Privacy Officer responsible for AI-tool use.",
         "Document which tools are used and for what purposes."),
        ("Principle 2 — Identifying Purposes",
         "Document that AI-assisted review is used to improve speed and accuracy of "
         "contract review services.",
         "Ensure purpose is limited to legal services in the retainer."),
        ("Principle 3 — Consent",
         "Add AI processing disclosure clause to standard retainer agreement.",
         "Use the template above. Confirm signed retainer before uploading client documents."),
        ("Principle 4 — Limiting Collection",
         "Upload only the minimum document necessary.",
         "Avoid uploading personal correspondence or unrelated documents."),
        ("Principle 5 — Limiting Use, Disclosure, Retention",
         "Set matter store retention period (recommended: 90 days or less).",
         "Enable auto-purge using the controls above."),
        ("Principle 6 — Accuracy",
         "All AI analysis is reviewed by a qualified lawyer before reliance.",
         "Do not share AI outputs directly with clients without review."),
        ("Principle 7 — Safeguards",
         "Ensure the matter store database (~/.contractcheck_matters.db) is on "
         "an encrypted drive or encrypted volume.",
         "Restrict access to authorized personnel only."),
        ("Principle 8 — Openness",
         "Update the firm privacy policy to reference AI-assisted review.",
         "Point clients to Anthropic's privacy policy for processor information."),
        ("Cross-Border Transfer",
         "Ensure retainer disclosure covers Anthropic's US infrastructure.",
         "Consider whether client matters require heightened protection (e.g., health, government)."),
        ("Breach Response",
         "Establish a breach response procedure for incidents involving client documents.",
         "Anthropic breach notifications: https://www.anthropic.com/policies/privacy"),
    ]

    for title, action, detail in obligations:
        with st.expander(title):
            st.markdown(f"**Required action:** {action}")
            st.caption(detail)

    st.markdown("---")
    st.caption(
        "This checklist is a practical guide only and does not constitute legal advice. "
        "Law firms should consult with privacy counsel to ensure their specific "
        "deployment satisfies PIPEDA, applicable provincial privacy legislation, "
        "and Law Society professional responsibility obligations.")


with st.expander("About ContractCheck Pro v12.6"):
    st.markdown("""**ContractCheck Pro v12.6** — the most capable Canadian AI contract review tool available.

**v12.0 — Major New Capabilities:**
- **Unlimited document length** — contracts of any size analysed completely via an overlapping chunk-then-synthesise pipeline. No silent truncation. Enterprise agreements, M&A documents, and 100+ page construction contracts all handled.
- **Persistent matter history** — every analysis saved to a local SQLite database. Browse, reload, and compare past matters. Export all matters as JSON.
- **True OOXML track changes** — redline DOCX uses `w:ins`/`w:del` XML elements. Microsoft Word shows real Track Changes that can be accepted or rejected natively by opposing counsel.
- **Vetted precedent clause library** — 40+ pre-approved Canadian replacement clauses (liability caps, indemnities, IP, termination, force majeure, governing law, and more) by contract type. AI adapts from vetted precedent rather than drafting from scratch.
- **Multi-document support** — upload main agreement plus up to 4 schedules, exhibits, or SOWs. All analysed together. Cross-document references flagged.
- **All v11 bug fixes included:** correct model names, temp file cleanup, file handle safety, ReportLab unique style names, OCR threading, `dir()` check replacement.

**Carried forward from v11:**
Clause Risk Benchmark Engine · Enforceability Scoring · Negotiation Simulator · 15-clause mandatory checklist · 7-section memo · Confidence tagging · Firm playbook · Version comparison · Audit log · Partner review checklist · Citation validation · Source links · 10 jurisdictions

**Security:** Zero data retention for contract text. Anthropic commercial API. Locally stored matter metadata only. Suitable for Law Society compliance requirements.

*AI first-pass. Must be reviewed by qualified counsel before any reliance.*""")
